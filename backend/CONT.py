import json
import subprocess
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import mysql.connector
import datetime
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.dialogs import Messagebox
from ttkbootstrap.tableview import Tableview
from ttkbootstrap.validation import add_regex_validation
from ttkbootstrap.widgets import DateEntry
from docx.shared import RGBColor
import zipfile
from io import BytesIO
import os
import winsound
import threading
import tempfile
from plyer import notification
import time
import hashlib
import time
import schedule

class BackgroundAlertService:
    def __init__(self, contrat_app):
        self.contrat_app = contrat_app  # Référence à ton application principale
        self.contract_hash = {}  # Pour stocker les contrats déjà notifiés
        self.sound_file = os.path.normpath(r"D:\UIAlert_Notification lasolisa 4 (ID 2066)_LS.wav")
        self.sound_enabled = True
        self.running = True
        self.hash_file = "notified_contracts.json"
        self.load_notified_contracts()

    def load_notified_contracts(self):
        """Charge les contrats déjà notifiés depuis un fichier JSON."""
        try:
            if os.path.exists(self.hash_file):
                with open(self.hash_file, 'r') as f:
                    self.contract_hash = json.load(f)
        except Exception as e:
            print(f"Erreur lors du chargement des contrats notifiés : {e}")

    def save_notified_contracts(self):
        """Sauvegarde les contrats notifiés dans un fichier JSON."""
        try:
            with open(self.hash_file, 'w') as f:
                json.dump(self.contract_hash, f)
        except Exception as e:
            print(f"Erreur lors de la sauvegarde des contrats notifiés : {e}")

    def check_expiring_contracts(self):
        """Vérifie les contrats CDD expirant dans les 30 jours."""
        try:
            today = datetime.datetime.now().date()
            with self.contrat_app.conn.cursor() as cursor:
                cursor.execute('''
                    SELECT e.matricule, e.nom, e.prenom, c.date_fin,
                           DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) AS jours_restants
                    FROM employes e
                    JOIN contrats c ON e.matricule = c.matricule
                    WHERE c.type_contrat = 'CDD'
                      AND c.date_fin IS NOT NULL
                      AND DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) BETWEEN 0 AND 30
                    ORDER BY jours_restants
                ''', (today, today))
                expiring_contracts = cursor.fetchall()

                for contract in expiring_contracts:
                    matricule, nom, prenom, date_fin, jours_restants = contract
                    contract_key = f"{matricule}_{date_fin.strip().lower()}"
                    contract_data = f"{matricule}{date_fin}"
                    contract_hash = hashlib.md5(contract_data.encode()).hexdigest()

                    if contract_key not in self.contract_hash or self.contract_hash[contract_key] != contract_hash:
                        # Nouvelle alerte
                        self.contract_hash[contract_key] = contract_hash
                        title = "Alerte Contrat"
                        message = f"Contrat de {nom} {prenom} (Matricule: {matricule}) expire le {date_fin} (dans {jours_restants} jours)"
                        try:
                            notification.notify(
                                title=title,
                                message=message,
                                app_name="Gestion des Contrats",
                                timeout=10
                            )
                            if self.sound_enabled and os.path.exists(self.sound_file):
                                os.system(f'start "" "{self.sound_file}"')  # pour Windows
                        except Exception as e:
                            print(f"Erreur lors de l'envoi de la notification : {e}")
                        self.save_notified_contracts()

        except Exception as e:
            print(f"Erreur dans la vérification des contrats : {e}")

    def run(self):
        """Exécute la vérification périodique des contrats."""
        print("Service d'arrière-plan démarré. Vérification toutes les 2 heures.")
        schedule.every(2).hours.do(self.check_expiring_contracts)
        while self.running:
            schedule.run_pending()
            time.sleep(60)  # Vérifie chaque minute si un job est planifié

    def stop(self):
        """Arrête le service d'arrière-plan."""
        self.running = False

class ContratApplication:

    def __init__(self, root):
        self.root = root
        self.root.title("Gestion des Contrats - Imbert Mnif")
        self.root.geometry("1600x900")
        self.style = ttk.Style(theme='flatly')
        self.style.configure("Treeview", rowheight=30, font=('Segoe UI', 10))
        self.style.configure("Treeview.Heading", font=('Segoe UI', 11, 'bold'))
        self.last_contract_check = None
        self.alert_timer = None
        self.check_interval = 4* 60 * 1000  # 2 minutes in milliseconds
        self.contract_hash = {}  # Store hashes of contract data to detect changes
        self.sound_file = os.path.normpath(r"D:\UIAlert_Notification lasolisa 4 (ID 2066)_LS.wav")
        if not os.path.exists(self.sound_file):
            Messagebox.show_warning(
                f"Le fichier sonore d'alerte n'a pas été trouvé à l'emplacement :\n{self.sound_file}\n\nLes alertes seront silencieuses.",
                "Avertissement",
                parent=root
            )

        # Configuration de la connexion à la base de données
        self.db_config = {
            "host": "192.168.1.210",  # Verify this is correct
            "user": "omar",  # Verify username
            "password": "1234",  # Verify password
            "database": "rh",  # Verify database exists
            "charset": "utf8mb4",
            "collation": "utf8mb4_unicode_ci",
            'port':"3306"# Default MySQL port

        }

        try:
            self.conn = mysql.connector.connect(**self.db_config)
            self.create_database()

        except mysql.connector.Error as err:
            error_msg = f"Database connection error:\n"
            if err.errno == mysql.connector.errorcode.ER_ACCESS_DENIED_ERROR:
                error_msg += "- Invalid credentials\n"
            elif err.errno == mysql.connector.errorcode.ER_BAD_DB_ERROR:
                error_msg += "- Database does not exist\n"
            else:
                error_msg += f"- {str(err)}\n"

            error_msg += f"Please verify your database configuration:\n"
            error_msg += f"Host: {self.db_config['host']}\n"
            error_msg += f"User: {self.db_config['user']}\n"
            error_msg += f"Database: {self.db_config['database']}"

            Messagebox.show_error("Database Connection Failed", error_msg)
            self.root.quit()
            return
            # Lancer le service d'arrière-plan
        self.start_background_service()
        self.variables = {
            "genre": tk.StringVar(value="féminin"),
            "matricule": tk.StringVar(),
            "contract_type": tk.StringVar(value="CDD"),
            "salary_type": tk.StringVar(value="hourly")
        }
        self.current_employee = None
        self.logo_path = r"D:\imbertlogo.png"
        self.undo_stack = []
        self.column_definitions = [
            {"text": "Matricule", "stretch": False, "width": 100},
            {"text": "Nom", "stretch": True, "width": 150},
            {"text": "Prénom", "stretch": True, "width": 150},
            {"text": "Genre", "stretch": False, "width": 80},
            {"text": "Date Naissance", "stretch": True, "width": 120},
            {"text": "Lieu Naissance", "stretch": True, "width": 150},
            {"text": "Adresse", "stretch": True, "width": 200},
            {"text": "Ville", "stretch": True, "width": 100},
            {"text": "CIN", "stretch": True, "width": 100},
            {"text": "Date CIN", "stretch": True, "width": 120},
            {"text": "Lieu CIN", "stretch": True, "width": 150},
            {"text": "Poste", "stretch": True, "width": 150},
            {"text": "Email", "stretch": True, "width": 200},
            {"text": "Téléphone", "stretch": True, "width": 120},
            {"text": "Type Contrat", "stretch": True, "width": 100},
            {"text": "Date Début", "stretch": True, "width": 120},
            {"text": "Date Fin", "stretch": True, "width": 120},
            {"text": "Salaire Base", "stretch": True, "width": 100},
            {"text": "Prime", "stretch": True, "width": 100},
            {"text": "Type Salaire", "stretch": True, "width": 100},
        ]

        self.setup_ui()
        self.load_data()
        # Lancer le service d'alerte en arrière-plan
        self.background_alert_service()
    def on_enter(self):
        self.setup_ui()
    CDD_MASCULIN = """
    طبقًا لأحكام الفصل 6-4 الجديد من مجلة الشغل

    تمهيد:حيث أن شركة أمبار منيف مختصة في صناعة الأحذية للتصدير"Sous Traitance" لفائدة حرفاء بالخارج تفوق في بعض الأحيان حجم قدرتها الإنتاجية العادية، ولما كانت هذه الطلبيات محدودة وغير منتظمة وغير مضمونة الاستمرارية فإن شركة أمبار منيف بحاجة إلى انتداب أجراء لمدة معينة لتلبية حاجياتها الظرفية نظرا للزيادة الغير عادية في حجم العمل المتأتية من ارتفاع طلبيات الحرفاء: Décathlon, Imac , Ricker et Cleon  . يعتبر الطرفان هذا التمهيد جزء من العقد ويلتزمان به كسائر فصوله.
    الحمد لله،
    -  عملا بأحكام مجلة الشغل ومقتضيات الاتفاقية  المشتركة القطاعية لصناعة الأحذية وتوابعها،
    - وبناء على طلب   {{Titre}} {{Prénom}} {{Nom}}  وتصريحه بأنه حر من كل التزام وغير مرتبط بعلاقة شغلية مع أي  مؤجر كان،
    تم الاتفاق والتراضي والتعاقد  بين الممضين أسفله
    أولا : شركة أمبار منيف، شركة خفية الاسم، مقرها الاجتماعي طريق قابس 3060 المحرس، ممثلة في شخص السيد ياسين المنيف بصفته مديرا عاما، المؤجر، من جهة
    ثانيا :  {{Titre}} {{Prénom}} {{Nom}} ، تاريخ الولادة {{DNAIS}}. مكانها {{LNAIS}} عنوان  {{Ville}} صاحب بطاقة التعريف القومية عدد  {{NCIN}} الصادرة بتاريخ {{DCIN}} بـ {{LCIN}}، بصفته متعاقد لمدة معينة، من جهة أخرى
    الفصل الأول : انتدبت شركة أمبار منيف بمقتضى هذا العقد {{Titre}} {{Prénom}} {{Nom}}  ليقع تشغيله بصفة {{Poste}} وذلك بمصنعها الكائن بالعنوان المذكور أو بجميع المصالح التابعة لها.
    الفصل الثاني : مدة العمل وطبيعته : اتفق الطرفان على أن تبدأ العلاقة الشغلية بتاريخ {{DPERIODE}} وتنتهي  بتاريخ  {{FPERIODE}}، وذلك لمدة محدودة قدرها {{DUREE}}.
    الفصل الثالث : يتقاضى {{Titre}} {{Prénom}} {{Nom}} مقابل انجاز العمل موضوع هذا العقد أجرا خاما خاضعا للتنقيصات الاجتماعية والجنائية تفصيله كالآتي :
      أجر أساسي: {{SBASE}} دينار في  {{MPAIE}}
      منح مختلفة: {{PRIME}} دينارًا عن الشهر الكامل
    الفصل الرابع : يعترف الأجير أنه اطلع على تراتيب العمل السارية داخل المؤسسة وتعهد باحترامها بصفته جزءا لا يتجزأ عن هذا العقد بما في ذلك أوقات العمل، كما يتعهد بالعناية بعملها والمحافظة على ممتلكات المؤسسة من مواد أولية وأدوات وآلات وغيرها، كما أن عليها قبول أي إحالة إلى مركز عمل أو مصلحة يقترحها عليها المؤجر.
    الفصل الخامس : ينتهي هذا العقد بانتهاء أجله المذكور إعلاه بدون سابق إعلام ويصبح الطرفان في حل من العلاقة الشغلية التي تربطهما. كما ينتهي هذا العقد باتفاق الطرفين أو عند انتهاء الأشغال موضوع هذا العقد أو عند قيام الأجير بهفوة فادحة أو ضعف في مؤهلاته أو تدني إنتاجه.
    الفصل السادس : يتعهد  {{Titre}} {{Prénom}} {{Nom}} باحترام سمعة المؤسسة والإمتناع عن إفشاء الأسرار المهنية والصناعية مما ينجر عنه الإضرار بمصالح المؤسسة حتى بعد انتهاء العقد. ويتحمل الأجير مسؤولية كل ضرر ينجر عن مخالفة هذا الشرط ويتعهد بأداء الغرامات اللازمة لجبره بدون مساس بالتتبعات الجزائية طبقا للقانون الجاري به العمل.
    الفصل السابع : يعترف الطرفان أنهما اطلعا على هذا العقد وأمضيا في حالة جواز ومعرفة، ويعينان مقر مخابرتهما بعنوانهما المذكور أعلاه كما يلتزم الأجير بإعلام مؤجره كتابيا بكل تغيير يطرأ عن البيانات الشخصية التي أدل بها بما في ذلك عنوانه ومقر سكناه في ظرف 48 ساعة من تاريخ التغيير.
    حرر بالمحرس في نظيرين بتاريخ {{DATE_CONTRAT}}
      	إمضاء المؤجر وختمه                                                            	 	     إمضاء الأجير معرف ب
    """

    CDD_FEMININ = """
    طبقًا لأحكام الفصل 6-4 الجديد من مجلة الشغل 

    تمهيد:حيث أن شركة أمبار منيف مختصة في صناعة الأحذية للتصدير"Sous Traitance" لفائدة حرفاء بالخارج تفوق في بعض الأحيان حجم قدرتها الإنتاجية العادية، ولما كانت هذه الطلبيات محدودة وغير منتظمة وغير مضمونة الاستمرارية فإن شركة أمبار منيف بحاجة إلى انتداب أجراء لمدة معينة لتلبية حاجياتها الظرفية نظرا للزيادة الغير عادية في حجم العمل المتأتية من ارتفاع طلبيات الحرفاء: Décathlon, Imac , Ricker et Cleon  . يعتبر الطرفان هذا التمهيد جزء من العقد ويلتزمان به كسائر فصوله.
    الحمد لله،
    -  عملا بأحكام مجلة الشغل ومقتضيات الاتفاقية  المشتركة القطاعية لصناعة الأحذية وتوابعها،
    - وبناء على طلب   {{Titre}} {{Prénom}} {{Nom}}  وتصريحها بأنها حرة من كل التزام وغير مرتبطة بعلاقة شغلية مع أي  مؤجر كان،
    تم الاتفاق والتراضي والتعاقد  بين الممضين أسفله
    أولا : شركة أمبار منيف، شركة خفية الاسم، مقرها الاجتماعي طريق قابس 3060 المحرس، ممثلة في شخص السيد ياسين المنيف بصفته مديرا عاما، المؤجر، من جهة
    ثانيا :  {{Titre}} {{Prénom}} {{Nom}} ، تاريخ الولادة {{DNAIS}}. مكانها {{LNAIS}} عنوانها {{Ville}} صاحبة بطاقة التعريف القومية عدد  {{NCIN}} الصادرة بتاريخ {{DCIN}} بـ {{LCIN}}، بصفتها {{Poste}} متعاقدة لمدة معينة، من جهة أخرى
    الفصل الأول : انتدبت شركة أمبار منيف بمقتضى هذا العقد {{Titre}} {{Prénom}} {{Nom}}  ليقع تشغيلها بصفة  {{Poste}} وذلك بمصنعها الكائن بالعنوان المذكور أو بجميع المصالح التابعة لها.
    الفصل الثاني : مدة العمل وطبيعته : اتفق الطرفان على أن تبدأ العلاقة الشغلية بتاريخ {{DPERIODE}} وتنتهي  بتاريخ  {{FPERIODE}}، وذلك لمدة محدودة قدرها {{DUREE}}.
    الفصل الثالث : تتقاضى {{Titre}} {{Prénom}} {{Nom}} مقابل انجاز العمل موضوع هذا العقد أجرا خاما خاضعا للتنقيصات الاجتماعية والجنائية تفصيله كالآتي :
      أجر أساسي: {{SBASE}} دينار في {{MPAIE}}
      منح مختلفة: {{PRIME}} دينارًا عن الشهر الكامل
    الفصل الرابع : تعترف الأجيرة أنها اطلعت على تراتيب العمل السارية داخل المؤسسة وتعهدت باحترامها بصفتها جزءا لا يتجزأ عن هذا العقد بما في ذلك أوقات العمل، كما تتعهد بالعناية بعملها والمحافظة على ممتلكات المؤسسة من مواد أولية وأدوات وآلات وغيرها، كما أن عليها قبول أي إحالة إلى مركز عمل أو مصلحة يقترحها عليها المؤجر.
    الفصل الخامس : ينتهي هذا العقد بانتهاء أجله المذكور إعلاه بدون سابق إعلام ويصبح الطرفان في حل من العلاقة الشغلية التي تربطهما. كما ينتهي هذا العقد باتفاق الطرفين أو عند انتهاء الأشغال موضوع هذا العقد أو عند قيام الأجيرة بهفوة فادحة أو ضعف في مؤهلاتها أو تدني إنتاجها.
    الفصل السادس : تتعهد  {{Titre}} {{Prénom}} {{Nom}} باحترام سمعة المؤسسة والإمتناع عن إفشاء الأسرار المهنية والصناعية مما ينجر عنه الإضرار بمصالح المؤسسة حتى بعد انتهاء العقد. وتتحمل الأجيرة مسؤولية كل ضرر ينجر عن مخالفة هذا الشرط وتتعهد بأداء الغرامات اللازمة لجبره بدون مساس بالتتبعات الجزائية طبقا للقانون الجاري به العمل.
    الفصل السابع : يعترف الطرفان أنهما اطلعا على هذا العقد وأمضيا في حالة جواز ومعرفة، ويعينان مقر مخابرتهما بعنوانهما المذكور أعلاه كما تلتزم الأجيرة بإعلام مؤجرها كتابيا بكل تغيير يطرأ عن البيانات الشخصية التي أدلت بها بما في ذلك عنوانها ومقر سكناها في ظرف 48 ساعة من تاريخ التغيير.
    حرر بالمحرس في نظيرين بتاريخ {{DATE_CONTRAT}}
      	إمضاء المؤجر وختمه                                                            	 	     إمضاء الأجيرة معرف ب
    """

    CDI_MASCULIN = """
    الحمد لله،
    -  عملا بأحكام مجلة الشغل ومقتضيات الاتفاقية المشتركة القطاعية لصناعة الأحذية وتوابعها،
    وبناء على طلب  {{Titre}} {{Prénom}} {{Nom}} وتصريحه بأنه حر من كل التزام وغير مرتبط بعلاقة شغليه مع أي مؤجر كان،
    تم الاتفاق والتراضي والتعاقد بين الممضين أسفله
    أولا: شركة أمبار منيف، شركة خفية الاسم، مقرها الاجتماعي طريق قابس 3060 المحرس، ممثلة في شخص السيد ياسين المنيف بصفته مديرا عاما، المؤجر، من جهة
    ثانيا : السيد  {{Titre}} {{Prénom}} {{Nom}} تاريخ الولادة {{DNAIS}}. مكانها {{LNAIS}} عنوانه {{Ville}} صاحب بطاقة التعريف القومية عدد {{NCIN}} الصادرة بتاريخ {{DCIN}} بتونس بصفته أجير متعاقد لمدة غير معينة، من جهة أخرى

    الفصل الأول : انتدبت شركة أمبار منيف بمقتضى هذا العقد  {{Titre}} {{Prénom}} {{Nom}} ليقع تشغيله بصفة {{Poste}} وذلك بمصنعها الكائن بالعنوان المذكور أو بجميع المصالح التابعة لها.

    الفصل الثاني : مدة العمل وطبيعته : اتفق الطرفان على أن تبدأ العلاقة الشغلية لمدة غير محددة من تاريخ {{DPERIODE}} .
     يخضع العامل المذكور أعلاه الى فترة تجربة مدتها ستة أشهر قابلة للتجديد مرة واحدة ولنفس المدة ويمكن لاحد طرفي العقد انهاء العمل به قبل انقضاء فترة التجربة بعد اعلام الطرف الأخر باي وسيلة تترك اثرا كتابيا وذلك قبل خمسة عشر يوما من انهاء فترة التجربة.

    الفصل الثالث: يتقاضى  {{Titre}} {{Prénom}} {{Nom}} مقابل انجاز العمل موضوع هذا العقد أجرا خاما خاضعا للتنقيصات الاجتماعية والجنائية تفصيله كالآتي:
    أجر أســـاسي {{SBASE}}   دينار  في الساعة
    منـــح مختلفة {{PRIME}} دينارا  عن الشهر الكامل

    الفصل الرابع: يعترف الأجير أنه اطلع على تراتيب العمل السارية داخل المؤسسة وتعهد باحترامها بصفتها جزءا لا يتجزأ عن هذا العقد بما في ذلك أوقات العمل، كما يتعهد بالعناية بعمله والمحافظة على ممتلكات المؤسسة من مواد أولية وأدوات وآلات وغيرها، كما أن عليه قبول أي إحالة إلى مركز عمل أو مصلحة يقترحها عليه المؤجر. 

    الفصل السادس : يتعهد السيد  {{Titre}} {{Prénom}} {{Nom}} باحترام سمعة المؤسسة والامتناع عن إفشاء الأسرار المهنية والصناعية مما ينجر عنه الإضرار بمصالح المؤسسة حتى بعد انتهاء العقد. ويتحمل الأجير مسؤولية كل ضرر ينجر عن مخالفة هذا الشرط ويتعهد بأداء الغرامات اللازمة لجبره بدون مساس بالتتبعات الجزائية طبقا للقانون الجاري به العمل.

    الفصل السابع: يعترف الطرفان أنهما اطلعا على هذا العقد وأمضيا في حالة جواز ومعرفة، ويعينان مقر مخابرتهما بعنوانهما المذكور أعلاه كما تلتزم الأجيرة بإعلام مؤجرها كتابيا بكل تغيير يطرأ عن البيانات الشخصية التي أدلت بها بما في ذلك عنوانها ومقر سكناها في ظرف 48 ساعة من تاريخ التغيير.
    حرر بالمحرس في نظيرين بتاريخ {{DATE_CONTRAT}}
          إمضاء المؤجر وختمه                     			                    إمضاء الأجير معرف به
    """

    CDI_FEMININ = """
    الحمد لله،
    -  عملا بأحكام مجلة الشغل ومقتضيات الاتفاقية المشتركة القطاعية لصناعة الأحذية وتوابعها،
    وبناء على طلب  {{Titre}} {{Prénom}} {{Nom}} وتصريحها بأنها حرة من كل التزام وغير مرتبطة بعلاقة شغليه مع أي مؤجر كان،
    تم الاتفاق والتراضي والتعاقد بين الممضين أسفله
    أولا: شركة أمبار منيف، شركة خفية الاسم، مقرها الاجتماعي طريق قابس 3060 المحرس، ممثلة في شخص السيد ياسين المنيف بصفته مديرا عاما، المؤجر، من جهة
    ثانيا : الآنسة {{Titre}} {{Prénom}} {{Nom}} تاريخ الولادة {{DNAIS}}. مكانها {{LNAIS}} عنوانها {{Ville}} صاحبة بطاقة التعريف القومية عدد {{NCIN}} الصادرة بتاريخ {{DCIN}} بتونس بصفتها أجيرة متعاقدة لمدة غير معينة، من جهة أخرى

    الفصل الأول : انتدبت شركة أمبار منيف بمقتضى هذا العقد  {{Titre}} {{Prénom}} {{Nom}} ليقع تشغيلها بصفة {{Poste}} وذلك بمصنعها الكائن بالعنوان المذكور أو بجميع المصالح التابعة لها.

    الفصل الثاني : مدة العمل وطبيعته : اتفق الطرفان على أن تبدأ العلاقة الشغلية لمدة غير محددة من تاريخ {{DPERIODE}} .
     تخضع العاملة المذكورة أعلاه الى فترة تجربة مدتها ستة أشهر قابلة للتجديد مرة واحدة ولنفس المدة ويمكن لاحد طرفي العقد انهاء العمل به قبل انقضاء فترة التجربة بعد اعلام الطرف الأخر باي وسيلة تترك اثرا كتابيا وذلك قبل خمسة عشر يوما من انهاء فترة التجربة.

    الفصل الثالث: تتقاضى  {{Titre}} {{Prénom}} {{Nom}} مقابل انجاز العمل موضوع هذا العقد أجرا خاما خاضعا للتنقيصات الاجتماعية والجنائية تفصيله كالآتي:
    أجر أســـاسي {{SBASE}}   دينار  في الساعة
    منـــح مختلفة {{PRIME}} دينارا  عن الشهر الكامل

    الفصل الرابع: تعترف الأجيرة أنها اطلعت على تراتيب العمل السارية داخل المؤسسة وتعهدت باحترامها بصفتها جزءا لا يتجزأ عن هذا العقد بما في ذلك أوقات العمل، كما تتعهد بالعناية بعملها والمحافظة على ممتلكات المؤسسة من مواد أولية وأدوات وآلات وغيرها، كما أن عليها قبول أي إحالة إلى مركز عمل أو مصلحة يقترحها عليها المؤجر. 

    الفصل السادس : تتعهد الآنسة {{Titre}} {{Prénom}} {{Nom}} باحترام سمعة المؤسسة والامتناع عن إفشاء الأسرار المهنية والصناعية مما ينجر عنه الإضرار بمصالح المؤسسة حتى بعد انتهاء العقد. وتتحمل الأجيرة مسؤولية كل ضرر ينجر عن مخالفة هذا الشرط وتتعهد بأداء الغرامات اللازمة لجبره بدون مساس بالتتبعات الجزائية طبقا للقانون الجاري به العمل.

    الفصل السابع: يعترف الطرفان أنهما اطلعا على هذا العقد وأمضيا في حالة جواز ومعرفة، ويعينان مقر مخابرتهما بعنوانهما المذكور أعلاه كما تلتزم الأجيرة بإعلام مؤجرها كتابيا بكل تغيير يطرأ عن البيانات الشخصية التي أدلت بها بما في ذلك عنوانها ومقر سكناها في ظرف 48 ساعة من تاريخ التغيير.
    حرر بالمحرس في نظيرين بتاريخ {{DATE_CONTRAT}}
          إمضاء المؤجر وختمه                     			                    إمضاء الأجيرة معرف بها
    """

    def start_background_service(self):
        """Lance le service d'arrière-plan dans un thread séparé."""
        self.alert_service = BackgroundAlertService(self)
        self.background_thread = threading.Thread(target=self.alert_service.run, daemon=True)
        self.background_thread.start()
        print("Service d'arrière-plan démarré")



    def create_database(self):
        cursor = self.conn.cursor()
        cursor.execute(f'''
            CREATE TABLE IF NOT EXISTS employes (
                matricule VARCHAR(50) PRIMARY KEY,
                nom VARCHAR(100) NOT NULL,
                prenom VARCHAR(100) NOT NULL,
                genre VARCHAR(20) NOT NULL,
                date_naissance VARCHAR(10),
                lieu_naissance VARCHAR(100),
                adresse VARCHAR(200),
                ville VARCHAR(100) DEFAULT 'المحرس',
                cin VARCHAR(20),
                date_cin VARCHAR(10),
                lieu_cin VARCHAR(100) DEFAULT 'تونس',
                poste VARCHAR(100),
                email VARCHAR(100),
                telephone VARCHAR(20)
            ) CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci
        ''')
        cursor.execute(f'''
            CREATE TABLE IF NOT EXISTS contrats (
                id INT AUTO_INCREMENT PRIMARY KEY,
                matricule VARCHAR(50),
                type_contrat VARCHAR(10),
                date_creation VARCHAR(20),
                date_debut VARCHAR(10),
                date_fin VARCHAR(10),
                salaire_base FLOAT,
                prime FLOAT,
                salary_type VARCHAR(20),
                texte_contrat TEXT,
                FOREIGN KEY (matricule) REFERENCES employes(matricule)
            ) CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci
        ''')
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_matricule ON employes(matricule)")
        cursor.execute("CREATE INDEX IF NOT EXISTS idx_contrat_matricule ON contrats(matricule)")
        self.conn.commit()

    def setup_ui(self):
        main_panel = ttk.Frame(self.root)
        main_panel.pack(fill=BOTH, expand=True, padx=15, pady=15)

        self.notebook = ttk.Notebook(main_panel, bootstyle=PRIMARY)
        self.notebook.pack(fill=BOTH, expand=True)

        self.create_employee_tab(main_panel)
        self.create_search_tab(main_panel)
        self.create_contract_tab(main_panel)
        self.create_list_tab(main_panel)

        # Créer le cadre de statut avant de l'utiliser
        status_frame = ttk.Frame(main_panel, bootstyle=INFO)
        status_frame.pack(fill=X, pady=(10, 0))

        self.status_var = tk.StringVar(value="Prêt")
        ttk.Label(status_frame, textvariable=self.status_var, bootstyle=(INFO, INVERSE),
                  font=('Segoe UI', 10)).pack(side=LEFT, padx=10)

        # Ajouter le bouton d'aide et à propos
        ttk.Button(status_frame, text="Aide", command=self.show_help,
                   bootstyle=(INFO, OUTLINE)).pack(side=RIGHT, padx=5)
        ttk.Button(status_frame, text="À propos", command=self.show_about,
                   bootstyle=(INFO, OUTLINE)).pack(side=RIGHT, padx=5)

        # Option pour les alertes sonores (si vous voulez garder cette fonctionnalité)
        self.sound_enabled = True
        ttk.Checkbutton(
            status_frame,
            text="Alertes sonores",
            variable=tk.BooleanVar(value=self.sound_enabled),
            command=lambda: setattr(self, 'sound_enabled', not self.sound_enabled),
            bootstyle="round-toggle"
        ).pack(side=RIGHT, padx=5)

    def create_employee_tab(self, parent):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Nouvel Employé")

        # Create a canvas and scrollbar
        canvas = tk.Canvas(frame, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview, bootstyle=PRIMARY)
        scrollable_frame = ttk.Frame(canvas)

        # Configure canvas
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Pack canvas and scrollbar
        canvas.pack(side=LEFT, fill=BOTH, expand=True, padx=10, pady=10)
        scrollbar.pack(side=RIGHT, fill=Y)

        fields = [
            ("Matricule*", "matricule", r'^\w+$', ttk.Entry),
            ("Nom*", "nom", None, ttk.Entry),
            ("Prénom*", "prenom", None, ttk.Entry),
            ("Date Naissance (JJ/MM/AAAA)", "date_naissance", r'^\d{2}/\d{2}/\d{4}$', DateEntry),
            ("Lieu Naissance", "lieu_naissance", None, ttk.Entry),
            ("Adresse", "adresse", None, ttk.Entry),
            ("Ville", "ville", None, ttk.Entry),
            ("CIN", "cin", r'^\d{8}$', ttk.Entry),
            ("Date CIN (JJ/MM/AAAA)", "date_cin", r'^\d{2}/\d{2}/\d{4}$', DateEntry),
            ("Lieu CIN", "lieu_cin", None, ttk.Entry),
            ("Poste", "poste", None, ttk.Entry),
            ("Email", "email", r'^[^@]+@[^@]+\.[^@]+$', ttk.Entry),
            ("Téléphone", "telephone", r'^\+?\d{10,12}$', ttk.Entry)
        ]

        self.entries = {}
        form_frame = ttk.LabelFrame(scrollable_frame, text="Informations Employé", bootstyle=PRIMARY)
        form_frame.pack(fill=BOTH, padx=10, pady=10, expand=True)

        for i, (label, field, regex, widget_type) in enumerate(fields):
            ttk.Label(form_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=E)
            entry = widget_type(form_frame, bootstyle="primary") if widget_type != DateEntry else DateEntry(form_frame,
                                                                                                            bootstyle="primary",
                                                                                                            dateformat="%d/%m/%Y")
            if widget_type == DateEntry:
                entry.entry.configure(justify="center")  # Center date text for clarity
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=EW)
            if regex and widget_type == ttk.Entry:
                add_regex_validation(entry, regex)
                entry.bind("<KeyRelease>", lambda e, f=field: self.validate_field(e.widget, f))
            self.entries[field] = entry

        # Add Gender Selection
        ttk.Label(form_frame, text="Genre*", font=('Segoe UI', 10)).grid(row=len(fields), column=0, sticky=E, padx=5,
                                                                         pady=5)
        genre_frame = ttk.Frame(form_frame)
        genre_frame.grid(row=len(fields), column=1, sticky=W)
        ttk.Radiobutton(genre_frame, text="Féminin", variable=self.variables["genre"], value="féminin",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)
        ttk.Radiobutton(genre_frame, text="Masculin", variable=self.variables["genre"], value="masculin",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)

        self.entries['ville'].insert(0, "المحرس")
        self.entries['lieu_cin'].insert(0, "تونس")

        contract_frame = ttk.LabelFrame(scrollable_frame, text="Détails du Contrat", bootstyle=PRIMARY)
        contract_frame.pack(fill=BOTH, padx=10, pady=10, expand=True)

        ttk.Label(contract_frame, text="Type de Contrat*", font=('Segoe UI', 10)).grid(row=0, column=0, padx=5, pady=5,
                                                                                       sticky=E)
        ttk.Radiobutton(contract_frame, text="CDD", variable=self.variables["contract_type"], value="CDD",
                        bootstyle="primary-toolbutton", command=self.toggle_date_fin).grid(row=0, column=1, sticky=W)
        ttk.Radiobutton(contract_frame, text="CDI", variable=self.variables["contract_type"], value="CDI",
                        bootstyle="primary-toolbutton", command=self.toggle_date_fin).grid(row=0, column=2, sticky=W)

        contract_fields = [
            ("Date Début (JJ/MM/AAAA)*", "date_debut", r'^\d{2}/\d{2}/\d{4}$', DateEntry),
            ("Date Fin (JJ/MM/AAAA)", "date_fin", r'^\d{2}/\d{2}/\d{4}$', DateEntry),
            ("Salaire Base*", "salaire", r'^\d+(\.\d{1,2})?$', ttk.Entry),
            ("Prime*", "prime", r'^\d+(\.\d{1,2})?$', ttk.Entry)
        ]

        self.contract_entries = {}
        for i, (label, field, regex, widget_type) in enumerate(contract_fields):
            ttk.Label(contract_frame, text=label, font=('Segoe UI', 10)).grid(row=i + 1, column=0, padx=5, pady=5,
                                                                              sticky=E)
            entry = widget_type(contract_frame, bootstyle="primary") if widget_type != DateEntry else DateEntry(
                contract_frame, bootstyle="primary", dateformat="%d/%m/%Y")
            if widget_type == DateEntry:
                entry.entry.configure(justify="center")
                if field == "date_fin" and self.variables["contract_type"].get() == "CDI":
                    entry.entry.config(state=DISABLED)
            entry.grid(row=i + 1, column=1, padx=5, pady=5, sticky=EW)
            if regex and widget_type == ttk.Entry:
                add_regex_validation(entry, regex)
                entry.bind("<KeyRelease>", lambda e, f=field: self.validate_field(e.widget, f))
            self.contract_entries[field] = entry

        # Set default values for date fields
        self.contract_entries['date_debut'].entry.delete(0, tk.END)
        self.contract_entries['date_debut'].entry.insert(0, datetime.datetime.now().strftime("%d/%m/%Y"))
        self.contract_entries['salaire'].insert(0, "")
        self.contract_entries['prime'].insert(0, "")

        ttk.Label(contract_frame, text="Type de Salaire*", font=('Segoe UI', 10)).grid(row=len(contract_fields) + 1,
                                                                                       column=0, padx=5, pady=5,
                                                                                       sticky=E)
        salary_type_frame = ttk.Frame(contract_frame)
        salary_type_frame.grid(row=len(contract_fields) + 1, column=1, sticky=W)
        ttk.Radiobutton(salary_type_frame, text="Par Heure", variable=self.variables["salary_type"], value="hourly",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)
        ttk.Radiobutton(salary_type_frame, text="Par Mois", variable=self.variables["salary_type"], value="monthly",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)

        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(fill=X, pady=10)
        ttk.Button(button_frame, text="Enregistrer et Générer", command=self.save_and_generate, bootstyle=SUCCESS).pack(
            side=LEFT, padx=5)
        ttk.Button(button_frame, text="Réinitialiser", command=self.clear_form, bootstyle=WARNING).pack(side=LEFT,
                                                                                                        padx=5)

        scrollable_frame.columnconfigure(1, weight=1)

        # Enable mouse wheel scrolling
        def on_mousewheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        canvas.bind_all("<MouseWheel>", on_mousewheel)

    def create_search_tab(self, parent):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Recherche")

        search_frame = ttk.Frame(frame)
        search_frame.pack(fill=X, padx=10, pady=10)
        ttk.Label(search_frame, text="Matricule:", font=('Segoe UI', 10)).pack(side=LEFT)
        self.search_combo = ttk.Combobox(search_frame, textvariable=self.variables["matricule"], font=('Segoe UI', 10))
        self.search_combo.pack(side=LEFT, padx=5, expand=True, fill=X)
        ttk.Button(search_frame, text="Rechercher", command=self.search_employee, bootstyle=INFO).pack(side=LEFT, padx=5)

        info_frame = ttk.LabelFrame(frame, text="Informations Employé", bootstyle=PRIMARY)
        info_frame.pack(fill=BOTH, expand=True, padx=10, pady=5)
        self.info_text = tk.Text(info_frame, wrap=WORD, height=12, font=('Segoe UI', 10))
        scrollbar = ttk.Scrollbar(info_frame, command=self.info_text.yview, bootstyle=PRIMARY)
        self.info_text.config(yscrollcommand=scrollbar.set)
        self.info_text.pack(side=LEFT, fill=BOTH, expand=True)
        scrollbar.pack(side=RIGHT, fill=Y)

        button_frame = ttk.Frame(frame)
        button_frame.pack(fill=X, padx=10, pady=10)
        self.view_contract_btn = ttk.Button(button_frame, text="Voir Contrat", command=self.show_last_contract, bootstyle=(PRIMARY, OUTLINE), state=DISABLED)
        self.view_contract_btn.pack(side=LEFT, padx=5)
        self.edit_btn = ttk.Button(button_frame, text="Modifier Employé", command=lambda: self.edit_employee(self.current_employee['matricule']), bootstyle=(WARNING, OUTLINE), state=DISABLED)
        self.edit_btn.pack(side=LEFT, padx=5)
        self.delete_btn = ttk.Button(button_frame, text="Supprimer Employé", command=lambda: self.delete_employee(self.current_employee['matricule']), bootstyle=(DANGER, OUTLINE), state=DISABLED)
        self.delete_btn.pack(side=LEFT, padx=5)

    def create_contract_tab(self, parent):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Contrat")

        self.contract_text = tk.Text(frame, wrap=WORD, font=('Arial', 11))
        self.contract_text.pack(side=LEFT, fill=BOTH, expand=True)
        scrollbar = ttk.Scrollbar(frame, command=self.contract_text.yview, bootstyle=PRIMARY)
        self.contract_text.config(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=RIGHT, fill=Y)

        # Configure text widget for right-to-left appearance
        self.contract_text.tag_configure("rtl", justify="right")
        self.contract_text.insert(tk.END, "", "rtl")  # Apply RTL tag to text

        button_frame = ttk.Frame(frame)
        button_frame.pack(fill=X, padx=10, pady=10)
        ttk.Button(button_frame, text="Exporter Word", command=self.export_word, bootstyle=SUCCESS).pack(side=LEFT,
                                                                                                         padx=5)
        ttk.Button(button_frame, text="Copier", command=self.copy_contract, bootstyle=INFO).pack(side=LEFT, padx=5)
        ttk.Button(button_frame, text="Imprimer", command=self.print_contract, bootstyle=PRIMARY).pack(side=LEFT, padx=5)

    # Ajoutez la méthode d'impression
    def print_contract(self):
        """Imprime le contrat actuellement affiché et affiche un aperçu"""
        if not self.contract_text.get(1.0, tk.END).strip():
            Messagebox.show_warning("Aucun contrat à imprimer", "Attention")
            return

        try:
            # Create temporary Word document
            temp_dir = tempfile.gettempdir()
            temp_file = os.path.join(temp_dir, f"contrat_temp_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")

            # Generate Word document
            doc = self.create_contract_doc(
                self.current_employee['matricule'],
                self.contract_text.get(1.0, tk.END).strip()
            )
            doc.save(temp_file)

            try:
                # Try printing with win32print if available
                import win32print
                import win32api

                printer_name = win32print.GetDefaultPrinter()
                win32api.ShellExecute(
                    0,
                    "print",
                    temp_file,
                    f'/d:"{printer_name}"',
                    temp_dir,
                    0
                )
                self.status_var.set(f"Contrat envoyé à l'imprimante {printer_name}")

            except ImportError:
                # Fallback for systems without win32print
                if os.name == 'posix':
                    subprocess.run(['lpr', temp_file])
                    self.status_var.set("Contrat envoyé à l'imprimante par défaut")
                else:
                    os.startfile(temp_file, "print")
                    self.status_var.set("Ouverture du contrat pour impression")

            # Open the document for preview
            try:
                if os.name == 'nt':
                    os.startfile(temp_file)
                else:
                    subprocess.run(['xdg-open', temp_file])
                Messagebox.show_info(
                    f"Le contrat a été envoyé à l'imprimante et ouvert pour aperçu.\n"
                    f"Fichier temporaire: {temp_file}",
                    "Impression et Aperçu historically"
                )

            except Exception as preview_error:
                Messagebox.show_warning(
                    f"Contrat imprimé, mais erreur lors de l'ouverture de l'aperçu:\n{str(preview_error)}\n"
                    f"Fichier temporaire: {temp_file}",
                    "Avertissement"
                )

        except Exception as e:
            Messagebox.show_error(
                f"Erreur lors de la création/impression du contrat:\n{str(e)}",
                "Erreur d'impression"
            )

    def create_list_tab(self, parent):
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Liste Employés")

        # Frame pour les statistiques et actions
        summary_frame = ttk.Frame(frame, bootstyle=INFO)
        summary_frame.pack(fill=X, padx=10, pady=5)

        # Statistiques principales
        stats_frame = ttk.Frame(summary_frame)
        stats_frame.pack(side=LEFT, fill=X, expand=True)

        # Labels pour les statistiques
        self.total_label = ttk.Label(stats_frame, text="Total: 0", font=('Segoe UI', 10))
        self.total_label.pack(side=LEFT, padx=10)

        self.cdd_label = ttk.Label(stats_frame, text="CDD: 0", font=('Segoe UI', 10))
        self.cdd_label.pack(side=LEFT, padx=10)

        self.cdi_label = ttk.Label(stats_frame, text="CDI: 0", font=('Segoe UI', 10))
        self.cdi_label.pack(side=LEFT, padx=10)

        self.salary_label = ttk.Label(stats_frame, text="Salaire Moyen: 0.00 TND", font=('Segoe UI', 10))
        self.salary_label.pack(side=LEFT, padx=10)

        # Bouton d'actualisation
        refresh_btn = ttk.Button(
            summary_frame,
            text="🔄 Actualiser",
            command=self.load_employee_table,
            bootstyle=(INFO, OUTLINE)
        )
        refresh_btn.pack(side=RIGHT, padx=5)

        filter_frame = ttk.Frame(frame)
        filter_frame.pack(fill=X, padx=10, pady=5)
        ttk.Label(filter_frame, text="Filtrer par:", font=('Segoe UI', 10)).pack(side=LEFT)
        self.filter_var = tk.StringVar()
        self.filter_combo = ttk.Combobox(filter_frame, textvariable=self.filter_var,
                                         values=["Nom", "Matricule", "Type Contrat"],
                                         font=('Segoe UI', 10))
        self.filter_combo.pack(side=LEFT, padx=5)
        self.filter_entry = ttk.Entry(filter_frame, font=('Segoe UI', 10))
        self.filter_entry.pack(side=LEFT, padx=5, expand=True, fill=X)
        ttk.Button(filter_frame, text="Filtrer", command=self.apply_filter, bootstyle=INFO).pack(side=LEFT, padx=5)
        ttk.Button(filter_frame, text="Réinitialiser", command=self.reset_filter, bootstyle=WARNING).pack(side=LEFT,
                                                                                                          padx=5)
        ttk.Button(filter_frame, text="Exporter Tous", command=self.export_all_contracts, bootstyle=SUCCESS).pack(
            side=LEFT, padx=5)
        ttk.Button(filter_frame, text="Annuler", command=self.undo_action, bootstyle=(WARNING, OUTLINE)).pack(side=LEFT,
                                                                                                              padx=5)

        self.employee_table = Tableview(
            frame, coldata=self.column_definitions, rowdata=[], paginated=True, searchable=True,
            bootstyle=PRIMARY, autoalign=True, stripecolor=('lightblue', None), pagesize=20
        )
        self.employee_table.pack(fill=BOTH, expand=True, padx=10, pady=10)

        self.employee_table.view.bind("<Double-1>", self.edit_cell)
        self.context_menu = tk.Menu(self.root, tearoff=0, font=('Segoe UI', 10))
        self.context_menu.add_command(label="Modifier", command=self.context_menu_modify)
        self.context_menu.add_command(label="Supprimer", command=self.context_menu_delete)
        self.context_menu.add_command(label="Voir Contrat", command=self.context_menu_view_contract)
        self.context_menu.add_command(label="Exporter Contrat", command=self.context_menu_export_contract)
        self.employee_table.view.bind("<Button-3>", self.show_context_menu)


    def validate_field(self, widget, field):
        value = widget.get()
        validators = {
            "matricule": lambda v: bool(re.match(r'^\w+$', v)) if v else False,
            "cin": lambda v: bool(re.match(r'^\d{8}$', v)) if v else True,
            "date_naissance": lambda v: bool(re.match(r'^\d{2}/\d{2}/\d{4}$', v)) if v else True,
            "date_cin": lambda v: bool(re.match(r'^\d{2}/\d{2}/\d{4}$', v)) if v else True,
            "email": lambda v: bool(re.match(r'^[^@]+@[^@]+\.[^@]+$', v)) if v else True,
            "telephone": lambda v: bool(re.match(r'^\+?\d{10,12}$', v)) if v else True,
            "salaire": lambda v: bool(re.match(r'^\d+(\.\d{1,2})?$', v)) and float(v) > 0 if v else False,
            "prime": lambda v: bool(re.match(r'^\d+(\.\d{1,2})?$', v)) and float(v) >= 0 if v else False,
            "date_debut": lambda v: bool(re.match(r'^\d{2}/\d{2}/\d{4}$', v)) if v else False,
            "date_fin": lambda v: bool(re.match(r'^\d{2}/\d{2}/\d{4}$', v)) if v else True
        }
        widget.configure(bootstyle="danger" if not validators.get(field, lambda x: True)(value) else "primary")

    def toggle_date_fin(self):
        state = DISABLED if self.variables["contract_type"].get() == "CDI" else NORMAL
        self.contract_entries['date_fin'].entry.config(state=state)

    def load_data(self):
        self.load_matricules()
        self.load_employee_table()
        # Vérifier les contrats proches d'expiration après un court délai

    def load_matricules(self):
        with self.conn.cursor() as cursor:
            cursor.execute("SELECT matricule FROM employes ORDER BY matricule")
            self.search_combo['values'] = [row[0] for row in cursor.fetchall()]

    def load_employee_table(self):
        with self.conn.cursor() as cursor:
            cursor.execute('''
                SELECT e.matricule, e.nom, e.prenom, e.genre, e.date_naissance, e.lieu_naissance,
                       e.adresse, e.ville, e.cin, e.date_cin, e.lieu_cin, e.poste, e.email, e.telephone,
                       c.type_contrat, c.date_debut, c.date_fin, c.salaire_base, c.prime, c.salary_type
                FROM employes e
                LEFT JOIN contrats c ON e.matricule = c.matricule
                AND c.id = (SELECT MAX(id) FROM contrats WHERE matricule = e.matricule)
            ''')
            self.update_table_data(cursor.fetchall())
            self.status_var.set(f"{len(self.employee_table.get_rows())} employés chargés")

    def update_table_data(self, rows):
        self.employee_table.delete_rows()
        today = datetime.datetime.now().date()

        for row in rows:
            matricule = row[0]
            action_frame = ttk.Frame(self.employee_table.view)

            # Vérifier si le contrat expire bientôt
            warning = ""
            if row[16]:  # Si date_fin existe
                try:
                    end_date = datetime.datetime.strptime(row[16], "%d/%m/%Y").date()
                    days_left = (end_date - today).days
                    if 0 <= days_left <= 30:
                        warning = "⚠️ "  # Ajouter un emoji d'avertissement
                except ValueError:
                    pass

            ttk.Button(action_frame, text="Modifier",
                       command=lambda m=matricule: self.edit_employee(m),
                       bootstyle=(PRIMARY, OUTLINE), width=8).pack(side=LEFT, padx=2)
            ttk.Button(action_frame, text="Supprimer",
                       command=lambda m=matricule: self.delete_employee(m),
                       bootstyle=(DANGER, OUTLINE), width=8).pack(side=LEFT, padx=2)
            ttk.Button(action_frame, text="Contrat",
                       command=lambda m=matricule: self.view_contract_from_table(m),
                       bootstyle=(INFO, OUTLINE), width=8).pack(side=LEFT, padx=2)

            # Ajouter l'avertissement au nom si nécessaire
            display_row = list(row[:14])
            display_row[1] = warning + display_row[1]  # Ajouter l'avertissement au nom

            display_row += [
                row[14] or "N/A",
                row[15] or "N/A",
                row[16] or "N/A",
                str(row[17]) if row[17] is not None else "N/A",
                str(row[18]) if row[18] is not None else "N/A",
                row[19] or "N/A",
                action_frame
            ]
            self.employee_table.insert_row(values=display_row)

        self.update_summary()


    def apply_filter(self):
        filter_type = self.filter_var.get()
        filter_value = self.filter_entry.get().strip().lower()
        if not filter_type or not filter_value:
            self.load_employee_table()
            return

        query = '''
            SELECT e.matricule, e.nom, e.prenom, e.genre, e.date_naissance, e.lieu_naissance,
                   e.adresse, e.ville, e.cin, e.date_cin, e.lieu_cin, e.poste, e.email, e.telephone,
                   c.type_contrat, c.date_debut, c.date_fin, c.salaire_base, c.prime, c.salary_type
            FROM employes e
            LEFT JOIN contrats c ON e.matricule = c.matricule
            AND c.id = (SELECT MAX(id) FROM contrats WHERE matricule = e.matricule)
            WHERE {}
        '''
        params = (f"%{filter_value}%",) if filter_type != "Type Contrat" else (filter_value.upper(),)
        condition = {
            "Nom": "LOWER(e.nom) LIKE %s",
            "Matricule": "e.matricule LIKE %s",
            "Type Contrat": "c.type_contrat = %s"
        }.get(filter_type, "1=1")

        with self.conn.cursor() as cursor:
            cursor.execute(query.format(condition), params)
            self.update_table_data(cursor.fetchall())
            self.status_var.set(f"{len(self.employee_table.get_rows())} employés trouvés")

    def reset_filter(self):
        self.filter_var.set("")
        self.filter_entry.delete(0, tk.END)
        self.load_employee_table()

    def show_context_menu(self, event):
        row_id = self.employee_table.view.identify_row(event.y)
        if row_id:
            self.employee_table.view.selection_set(row_id)
            self.selected_matricule = self.employee_table.get_row(row_id).values[0]
            self.context_menu.post(event.x_root, event.y_root)

    def context_menu_modify(self):
        self.edit_employee(self.selected_matricule)

    def context_menu_delete(self):
        self.delete_employee(self.selected_matricule)

    def context_menu_view_contract(self):
        self.current_employee = {'matricule': self.selected_matricule}
        self.show_last_contract()

    def context_menu_export_contract(self):
        self.current_employee = {'matricule': self.selected_matricule}
        self.show_last_contract()
        self.export_word()



    def clear_form(self):
        for entry in self.entries.values():
            if isinstance(entry, ttk.Entry):
                entry.delete(0, tk.END)
            elif isinstance(entry, DateEntry):  # Handle DateEntry differently
                entry.entry.delete(0, tk.END)

        for entry in self.contract_entries.values():
            if isinstance(entry, ttk.Entry):
                entry.delete(0, tk.END)
            elif isinstance(entry, DateEntry):  # Handle DateEntry differently
                entry.entry.delete(0, tk.END)
        self.entries['ville'].insert(0, "المحرس")
        self.entries['lieu_cin'].insert(0, "تونس")
        self.contract_entries['date_debut'].entry.insert(0, datetime.datetime.now().strftime("%d/%m/%Y"))
        self.contract_entries['salaire'].insert(0, "2500.00")
        self.contract_entries['prime'].insert(0, "500.00")
        self.variables["genre"].set("féminin")
        self.variables["contract_type"].set("CDD")
        self.variables["salary_type"].set("hourly")
        self.status_var.set("Formulaire réinitialisé")

    def search_employee(self):
        matricule = self.variables["matricule"].get()
        if not matricule:
            Messagebox.show_warning("Veuillez saisir un matricule", "Attention")
            return

        with self.conn.cursor() as cursor:
            cursor.execute('''
                SELECT e.*, c.type_contrat, c.date_debut, c.date_fin, c.salaire_base, c.prime, c.salary_type
                FROM employes e
                LEFT JOIN contrats c ON e.matricule = c.matricule
                AND c.id = (SELECT MAX(id) FROM contrats WHERE matricule = e.matricule)
                WHERE e.matricule = %s
            ''', (matricule,))
            employee = cursor.fetchone()

        if employee:
            self.current_employee = {
                'matricule': employee[0], 'nom': employee[1], 'prenom': employee[2], 'genre': employee[3],
                'date_naissance': employee[4], 'lieu_naissance': employee[5], 'adresse': employee[6],
                'ville': employee[7], 'cin': employee[8], 'date_cin': employee[9], 'lieu_cin': employee[10],
                'poste': employee[11], 'email': employee[12], 'telephone': employee[13],
                'type_contrat': employee[14] or "N/A", 'date_debut': employee[15] or "N/A",
                'date_fin': employee[16] or "N/A", 'salaire_base': str(employee[17]) if employee[17] is not None else "N/A",
                'prime': str(employee[18]) if employee[18] is not None else "N/A", 'salary_type': employee[19] or "N/A"
            }
            info = f"""Matricule: {employee[0]}
Nom: {employee[1]} {employee[2]}
Genre: {employee[3]}
Date Naissance: {employee[4] or 'N/A'} à {employee[5] or 'N/A'}
Adresse: {employee[6] or 'N/A'}, {employee[7] or 'N/A'}
CIN: {employee[8] or 'N/A'} (délivré le {employee[9] or 'N/A'} à {employee[10] or 'N/A'})
Poste: {employee[11] or 'N/A'}
Email: {employee[12] or 'N/A'}
Téléphone: {employee[13] or 'N/A'}
Type Contrat: {employee[14] or 'N/A'}
Date Début: {employee[15] or 'N/A'}
Date Fin: {employee[16] or 'N/A'}
Salaire Base: {employee[17] or 'N/A'} TND
Prime: {employee[18] or 'N/A'} TND
Type Salaire: {employee[19] or 'N/A'}"""
            self.info_text.delete(1.0, tk.END)
            self.info_text.insert(tk.END, info)
            self.view_contract_btn.config(state=NORMAL)
            self.edit_btn.config(state=NORMAL)
            self.delete_btn.config(state=NORMAL)
            self.status_var.set(f"Employé trouvé: {employee[1]} {employee[2]}")
        else:
            self.clear_search()

    def clear_search(self):
        self.info_text.delete(1.0, tk.END)
        self.current_employee = None
        self.view_contract_btn.config(state=DISABLED)
        self.edit_btn.config(state=DISABLED)
        self.delete_btn.config(state=DISABLED)
        Messagebox.show_info("Aucun employé trouvé", "Information")
        self.status_var.set("Aucun résultat")

    def show_last_contract(self):
        if not self.current_employee:
            Messagebox.show_warning("Veuillez rechercher un employé", "Attention")
            return

        with self.conn.cursor() as cursor:
            cursor.execute("""
                           SELECT texte_contrat, type_contrat, date_debut, date_fin, salaire_base, prime, salary_type
                           FROM contrats
                           WHERE matricule = %s
                           ORDER BY id DESC LIMIT 1
                           """, (self.current_employee['matricule'],))
            contract = cursor.fetchone()

        if contract:
            self.contract_text.delete(1.0, tk.END)
            self.contract_text.insert(tk.END, contract[0])
            self.variables["contract_type"].set(contract[1])
            self.variables["salary_type"].set(contract[6])

            # Correction pour DateEntry
            self.contract_entries['date_debut'].entry.delete(0, tk.END)
            self.contract_entries['date_debut'].entry.insert(0, contract[2] or "")

            self.contract_entries['date_fin'].entry.delete(0, tk.END)
            self.contract_entries['date_fin'].entry.insert(0, contract[3] or "")

            # Pour les Entry normaux
            self.contract_entries['salaire'].delete(0, tk.END)
            self.contract_entries['salaire'].insert(0, str(contract[4]) if contract[4] is not None else "2500.00")

            self.contract_entries['prime'].delete(0, tk.END)
            self.contract_entries['prime'].insert(0, str(contract[5]) if contract[5] is not None else "500.00")

            self.notebook.select(2)  # Sélectionne l'onglet Contrat
            self.status_var.set("Dernier contrat affiché")
        else:
            Messagebox.show_info("Aucun contrat trouvé", "Information")
            self.status_var.set("Aucun contrat")

    def view_contract_from_table(self, matricule):
        self.current_employee = {'matricule': matricule}
        self.show_last_contract()

    def export_word(self):
        """Exporte le contrat de l'employé sélectionné au format Word."""
        try:
            # Vérifier si un employé est sélectionné
            if not self.current_employee or not isinstance(self.current_employee, dict):
                Messagebox.show_error("Aucun employé sélectionné. Veuillez sélectionner un employé.", "Erreur")
                return

            # Générer le document Word avec les données de l'employé
            doc = self.create_contract_doc(self.current_employee['matricule'],
                                           self.current_employee.get('texte_contrat', ''))

            # Demander à l'utilisateur où sauvegarder le fichier
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Document Word", "*.docx")],
                initialfile=f"Contrat_{self.current_employee['matricule']}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            if not file_path:
                return  # L'utilisateur a annulé la sauvegarde

            # Sauvegarder le document
            doc.save(file_path)
            Messagebox.show_info(f"Contrat exporté avec succès sous {file_path}.", "Succès")
            self.status_var.set("Contrat exporté avec succès.")

        except Exception as e:
            Messagebox.show_error(f"Erreur lors de l'exportation du contrat : {str(e)}", "Erreur")
            print(f"Erreur dans export_word : {str(e)}")

    def export_all_contracts(self):
        with self.conn.cursor() as cursor:
            cursor.execute("SELECT matricule, texte_contrat FROM contrats WHERE id IN (SELECT MAX(id) FROM contrats GROUP BY matricule)")
            contracts = cursor.fetchall()

        if not contracts:
            Messagebox.show_warning("Aucun contrat à exporter", "Attention")
            return

        filepath = filedialog.asksaveasfilename(
            defaultextension=".zip", filetypes=[("Archive ZIP", "*.zip")],
            initialfile=f"Contrats_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
            title="Enregistrer tous les contrats"
        )
        if filepath:
            try:
                with zipfile.ZipFile(filepath, 'w', zipfile.ZIP_DEFLATED) as zipf:
                    for matricule, texte in contracts:
                        doc = self.create_contract_doc(matricule, texte)
                        doc_io = BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        zipf.writestr(f"Contrat_{matricule}.docx", doc_io.read())
                self.status_var.set(f"Exportation réussie: {filepath}")
                Messagebox.show_info(f"Tous les contrats exportés: {filepath}", "Succès")
            except Exception as e:
                Messagebox.show_error(f"Erreur d'exportation: {str(e)}", "Erreur")

    def copy_contract(self):
        self.root.clipboard_clear()
        self.root.clipboard_append(self.contract_text.get(1.0, tk.END))
        self.status_var.set("Texte du contrat copié")

    def edit_cell(self, event):
        row_id = self.employee_table.view.identify_row(event.y)
        column = self.employee_table.view.identify_column(event.x)
        if not row_id or not column:
            return

        col_idx = int(column.replace("#", "")) - 1
        col_name = self.column_definitions[col_idx]["text"]
        if col_name in ["Matricule", "Actions"]:
            return

        row_data = self.employee_table.get_row(row_id).values
        matricule = row_data[0]
        current_value = row_data[col_idx]

        entry = ttk.Entry(self.employee_table.view, bootstyle="primary", font=('Segoe UI', 10))
        entry.insert(0, current_value)
        entry.place(x=event.x_root - self.employee_table.view.winfo_rootx(), y=event.y_root - self.employee_table.view.winfo_rooty(), anchor="nw")

        def validate_input(value):
            validators = {
                "Date Naissance": r'^\d{2}/\d{2}/\d{4}$', "Date CIN": r'^\d{2}/\d{2}/\d{4}$',
                "Date Début": r'^\d{2}/\d{2}/\d{4}$', "Date Fin": r'^\d{2}/\d{2}/\d{4}$',
                "Email": r'^[^@]+@[^@]+\.[^@]+$', "Genre": r'^(féminin|masculin)$',
                "Type Contrat": r'^(CDD|CDI)$', "Salaire Base": r'^\d+(\.\d{1,2})?$',
                "Prime": r'^\d+(\.\d{1,2})?$', "Type Salaire": r'^(hourly|monthly)$'
            }
            return bool(re.match(validators.get(col_name, r'.*'), value)) and (float(value) > 0 if col_name in ["Salaire Base", "Prime"] and value else True)

        def save_edit(event=None):
            new_value = entry.get()
            if not validate_input(new_value):
                Messagebox.show_error(f"Valeur invalide pour {col_name}", "Erreur")
                entry.destroy()
                return

            try:
                with self.conn.cursor() as cursor:
                    sql_field = {
                        "Nom": "nom", "Prénom": "prenom", "Genre": "genre", "Date Naissance": "date_naissance",
                        "Lieu Naissance": "lieu_naissance", "Adresse": "adresse", "Ville": "ville",
                        "CIN": "cin", "Date CIN": "date_cin", "Lieu CIN": "lieu_cin", "Poste": "poste",
                        "Email": "email", "Téléphone": "telephone", "Type Contrat": "type_contrat",
                        "Date Début": "date_debut", "Date Fin": "date_fin", "Salaire Base": "salaire_base",
                        "Prime": "prime", "Type Salaire": "salary_type"
                    }.get(col_name)
                    if sql_field:
                        if col_name in ["Type Contrat", "Date Début", "Date Fin", "Salaire Base", "Prime", "Type Salaire"]:
                            cursor.execute("SELECT id FROM contrats WHERE matricule = %s ORDER BY id DESC LIMIT 1", (matricule,))
                            contract_id = cursor.fetchone()
                            if contract_id:
                                cursor.execute(f"UPDATE contrats SET {sql_field} = %s WHERE id = %s", (float(new_value) if col_name in ["Salaire Base", "Prime"] else new_value, contract_id[0]))
                                self.undo_stack.append(("contract_update", matricule, contract_id[0], sql_field, current_value))
                            else:
                                Messagebox.show_warning("Aucun contrat trouvé", "Attention")
                                entry.destroy()
                                return
                        else:
                            cursor.execute(f"UPDATE employes SET {sql_field} = %s WHERE matricule = %s", (new_value, matricule))
                            self.undo_stack.append(("employee_update", matricule, sql_field, current_value))
                        self.conn.commit()
                        self.load_employee_table()
                        self.status_var.set(f"Champ {col_name} mis à jour pour {matricule}")
                        if self.current_employee and self.current_employee['matricule'] == matricule:
                            self.search_employee()
            except Exception as e:
                Messagebox.show_error(f"Erreur de mise à jour: {str(e)}", "Erreur")
            finally:
                entry.destroy()

        entry.bind("<Return>", save_edit)
        entry.bind("<FocusOut>", save_edit)
        entry.focus_set()

    def validate_field_value(self, field, value):
        validators = {
            "matricule": lambda v: bool(re.match(r'^\w+$', v)) if v else False,
            "cin": lambda v: bool(re.match(r'^\d{8}$', v)) if v else True,
            "date_naissance": lambda v: self.is_valid_date(v) if v else True,
            "date_cin": lambda v: self.is_valid_date(v) if v else True,
            "email": lambda v: bool(re.match(r'^[^@]+@[^@]+\.[^@]+$', v)) if v else True,
            "telephone": lambda v: bool(re.match(r'^\+?\d{10,12}$', v)) if v else True,
            "salaire": lambda v: bool(re.match(r'^\d+(\.\d{1,2})?$', v)) and float(v) > 0 if v else False,
            "prime": lambda v: bool(re.match(r'^\d+(\.\d{1,2})?$', v)) and float(v) >= 0 if v else False,
            "date_debut": lambda v: self.is_valid_date(v) if v else False,
            "date_fin": lambda v: self.is_valid_date(v) if v else True
        }
        return validators.get(field, lambda x: True)(value)

    def is_valid_date(self, date_str):
        try:
            if not date_str or not re.match(r'^\d{2}/\d{2}/\d{4}$', date_str.strip()):
                return False
            datetime.datetime.strptime(date_str.strip(), "%d/%m/%Y")
            return True
        except ValueError:
            return False

    def save_employee_and_contract_changes(self, matricule, entries, contract_entries, genre_var, contract_type_var,
                                           salary_type_var, top):
        self.stop_alert_timer()
        self.check_expiring_contracts(force_notification=True)  # Change 'first_notification' to 'force_notification'
        try:
            # Validation des champs obligatoires
            required_employee = ['nom', 'prenom']
            required_contract = ['date_debut', 'salaire', 'prime']

            for field in required_employee:
                value = self.get_widget_value(entries[field])
                if not value:
                    Messagebox.show_error(f"Le champ {field} est obligatoire", "Erreur")
                    return

            for field in required_contract:
                value = self.get_widget_value(contract_entries[field])
                if not value:
                    Messagebox.show_error(f"Le champ {field} est obligatoire", "Erreur")
                    return

            # Validation des dates
            date_debut = self.get_widget_value(contract_entries['date_debut'])
            if not self.is_valid_date(date_debut):
                Messagebox.show_error("Format de date invalide pour la date de début (JJ/MM/AAAA)", "Erreur")
                return

            if contract_type_var.get() == "CDD":
                date_fin = self.get_widget_value(contract_entries['date_fin'])
                if not self.is_valid_date(date_fin):
                    Messagebox.show_error("Format de date invalide pour la date de fin (JJ/MM/AAAA)", "Erreur")
                    return

                if datetime.datetime.strptime(date_fin, "%d/%m/%Y") <= datetime.datetime.strptime(date_debut,
                                                                                                  "%d/%m/%Y"):
                    Messagebox.show_error("La date de fin doit être après la date de début", "Erreur")
                    return

            # Validation des valeurs numériques
            try:
                salaire = float(self.get_widget_value(contract_entries['salaire']))
                prime = float(self.get_widget_value(contract_entries['prime']))
            except ValueError:
                Messagebox.show_error("Le salaire et la prime doivent être des nombres valides", "Erreur")
                return

            with self.conn.cursor() as cursor:
                # Mise à jour employé
                cursor.execute('''
                               UPDATE employes
                               SET nom            = %s,
                                   prenom         = %s,
                                   genre          = %s,
                                   date_naissance = %s,
                                   lieu_naissance = %s,
                                   adresse        = %s,
                                   ville          = %s,
                                   cin            = %s,
                                   date_cin       = %s,
                                   lieu_cin       = %s,
                                   poste          = %s,
                                   email          = %s,
                                   telephone      = %s
                               WHERE matricule = %s
                               ''', (
                                   self.get_widget_value(entries['nom']),
                                   self.get_widget_value(entries['prenom']),
                                   genre_var.get(),
                                   self.get_widget_value(entries['date_naissance']) or None,
                                   self.get_widget_value(entries['lieu_naissance']) or None,
                                   self.get_widget_value(entries['adresse']) or None,
                                   self.get_widget_value(entries['ville']) or None,
                                   self.get_widget_value(entries['cin']) or None,
                                   self.get_widget_value(entries['date_cin']) or None,
                                   self.get_widget_value(entries['lieu_cin']) or None,
                                   self.get_widget_value(entries['poste']) or None,
                                   self.get_widget_value(entries['email']) or None,
                                   self.get_widget_value(entries['telephone']) or None,
                                   matricule
                               ))

                # Mise à jour contrat
                cursor.execute('''
                               UPDATE contrats
                               SET type_contrat = %s,
                                   date_debut   = %s,
                                   date_fin     = %s,
                                   salaire_base = %s,
                                   prime        = %s,
                                   salary_type  = %s
                               WHERE matricule = %s ORDER BY id DESC LIMIT 1
                               ''', (
                                   contract_type_var.get(),
                                   date_debut,
                                   self.get_widget_value(
                                       contract_entries['date_fin']) if contract_type_var.get() == "CDD" else None,
                                   salaire,
                                   prime,
                                   salary_type_var.get(),
                                   matricule
                               ))

                self.conn.commit()
                Messagebox.show_info("Mise à jour réussie", "Succès")
                self.load_data()
                top.destroy()

        except Exception as e:
            Messagebox.show_error(f"Erreur lors de la sauvegarde: {str(e)}", "Erreur")

    def edit_employee(self, matricule):
        top = ttk.Toplevel(self.root)
        top.title(f"Modifier Employé {matricule}")
        top.geometry("900x900")

        with self.conn.cursor() as cursor:
            cursor.execute("SELECT * FROM employes WHERE matricule = %s", (matricule,))
            employee = cursor.fetchone()
            cursor.execute(
                "SELECT type_contrat, date_debut, date_fin, salaire_base, prime, salary_type FROM contrats WHERE matricule = %s ORDER BY id DESC LIMIT 1",
                (matricule,))
            contract = cursor.fetchone()

        if not employee:
            Messagebox.show_error("Employé non trouvé", "Erreur")
            top.destroy()
            return

        notebook = ttk.Notebook(top, bootstyle=PRIMARY)
        notebook.pack(fill=BOTH, expand=True, padx=10, pady=10)

        # Frame pour les informations employé
        employee_frame = ttk.Frame(notebook)
        notebook.add(employee_frame, text="Détails Employé")

        # Frame pour les détails du contrat
        contract_frame = ttk.Frame(notebook)
        notebook.add(contract_frame, text="Détails Contrat")

        # Variables pour les radiobuttons
        genre_var = tk.StringVar(value=employee[3])
        contract_type_var = tk.StringVar(value=contract[0] if contract else "CDD")
        salary_type_var = tk.StringVar(value=contract[5] if contract else "hourly")

        # Dictionnaires pour stocker les références aux widgets
        entries = {}
        contract_entries = {}

        # Champs employé
        fields = [
            ("Matricule", "matricule", employee[0], True, ttk.Entry),
            ("Nom*", "nom", employee[1], False, ttk.Entry),
            ("Prénom*", "prenom", employee[2], False, ttk.Entry),
        ]

        for i, (label, field, value, disabled, widget_type) in enumerate(fields):
            ttk.Label(employee_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=E)
            entry = widget_type(employee_frame, bootstyle="primary")
            entry.insert(0, value)
            if disabled:
                entry.config(state='disabled')
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=EW)
            entries[field] = entry

        # Champs dates (utilisant DateEntry)
        date_fields = [
            ("Date Naissance (JJ/MM/AAAA)", "date_naissance", employee[4], False, employee_frame),
            ("Date CIN (JJ/MM/AAAA)", "date_cin", employee[9], False, employee_frame),
            ("Date Début (JJ/MM/AAAA)*", "date_debut", contract[1] if contract else "", False, contract_frame),
            ("Date Fin (JJ/MM/AAAA)", "date_fin", contract[2] if contract else "", contract_type_var.get() == "CDI",
             contract_frame),
        ]

        for i, (label, field, value, disabled, frame) in enumerate(date_fields):
            ttk.Label(frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=E)
            entry = DateEntry(frame, bootstyle="primary", dateformat="%d/%m/%Y")
            if value:
                entry.entry.delete(0, tk.END)
                entry.entry.insert(0, value)
            if disabled:
                entry.entry.config(state='disabled')
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=EW)

            if frame == employee_frame:
                entries[field] = entry
            else:
                contract_entries[field] = entry

        # Autres champs employé
        other_fields = [
            ("Lieu Naissance", "lieu_naissance", employee[5], False, ttk.Entry),
            ("Adresse", "adresse", employee[6], False, ttk.Entry),
            ("Ville", "ville", employee[7], False, ttk.Entry),
            ("CIN", "cin", employee[8], False, ttk.Entry),
            ("Lieu CIN", "lieu_cin", employee[10], False, ttk.Entry),
            ("Poste", "poste", employee[11], False, ttk.Entry),
            ("Email", "email", employee[12], False, ttk.Entry),
            ("Téléphone", "telephone", employee[13], False, ttk.Entry),
        ]

        start_row = len(date_fields) if frame == employee_frame else 0
        for i, (label, field, value, disabled, widget_type) in enumerate(other_fields, start=start_row):
            ttk.Label(employee_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=E)
            entry = widget_type(employee_frame, bootstyle="primary")
            entry.insert(0, value or "")
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=EW)
            entries[field] = entry

        # Genre
        ttk.Label(employee_frame, text="Genre*", font=('Segoe UI', 10)).grid(row=len(other_fields) + start_row,
                                                                             column=0, sticky=E, padx=5, pady=5)
        genre_frame = ttk.Frame(employee_frame)
        genre_frame.grid(row=len(other_fields) + start_row, column=1, sticky=W)
        ttk.Radiobutton(genre_frame, text="Féminin", variable=genre_var, value="féminin",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)
        ttk.Radiobutton(genre_frame, text="Masculin", variable=genre_var, value="masculin",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)

        # Type de contrat
        ttk.Label(contract_frame, text="Type de Contrat*", font=('Segoe UI', 10)).grid(row=len(date_fields), column=0,
                                                                                       padx=5, pady=5, sticky=E)
        contract_type_frame = ttk.Frame(contract_frame)
        contract_type_frame.grid(row=len(date_fields), column=1, sticky=W)
        ttk.Radiobutton(contract_type_frame, text="CDD", variable=contract_type_var, value="CDD",
                        bootstyle="primary-toolbutton",
                        command=lambda: contract_entries['date_fin'].entry.config(state=NORMAL)).pack(side=LEFT, padx=5)
        ttk.Radiobutton(contract_type_frame, text="CDI", variable=contract_type_var, value="CDI",
                        bootstyle="primary-toolbutton",
                        command=lambda: contract_entries['date_fin'].entry.config(state=DISABLED)).pack(side=LEFT,
                                                                                                        padx=5)

        # Salaire et prime
        salary_fields = [
            ("Salaire Base*", "salaire", str(contract[3]) if contract else "2500.00", False, ttk.Entry),
            ("Prime*", "prime", str(contract[4]) if contract else "500.00", False, ttk.Entry),
        ]

        for i, (label, field, value, disabled, widget_type) in enumerate(salary_fields, start=len(date_fields) + 1):
            ttk.Label(contract_frame, text=label, font=('Segoe UI', 10)).grid(row=i, column=0, padx=5, pady=5, sticky=E)
            entry = widget_type(contract_frame, bootstyle="primary")
            entry.insert(0, value)
            entry.grid(row=i, column=1, padx=5, pady=5, sticky=EW)
            contract_entries[field] = entry

        # Type de salaire
        ttk.Label(contract_frame, text="Type de Salaire*", font=('Segoe UI', 10)).grid(row=len(date_fields) + 3,
                                                                                       column=0, padx=5, pady=5,
                                                                                       sticky=E)
        salary_type_frame = ttk.Frame(contract_frame)
        salary_type_frame.grid(row=len(date_fields) + 3, column=1, sticky=W)
        ttk.Radiobutton(salary_type_frame, text="Par Heure", variable=salary_type_var, value="hourly",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)
        ttk.Radiobutton(salary_type_frame, text="Par Mois", variable=salary_type_var, value="monthly",
                        bootstyle="primary-toolbutton").pack(side=LEFT, padx=5)

        # Boutons
        button_frame = ttk.Frame(top)
        button_frame.pack(fill=X, pady=10)
        ttk.Button(button_frame, text="Enregistrer",
                   command=lambda: self.save_employee_and_contract_changes(
                       matricule, entries, contract_entries, genre_var, contract_type_var, salary_type_var, top),
                   bootstyle=SUCCESS).pack(side=LEFT, padx=5)
        ttk.Button(button_frame, text="Annuler", command=top.destroy, bootstyle=WARNING).pack(side=LEFT, padx=5)

        employee_frame.columnconfigure(1, weight=1)
        contract_frame.columnconfigure(1, weight=1)


    def delete_employee(self, matricule):
        if not Messagebox.yesno(f"Confirmer la suppression de l'employé {matricule}?", "Confirmation"):
            return

        try:
            with self.conn.cursor() as cursor:
                cursor.execute("SELECT * FROM employes WHERE matricule = %s", (matricule,))
                employee_data = cursor.fetchone()
                cursor.execute("SELECT * FROM contrats WHERE matricule = %s ORDER BY id DESC LIMIT 1", (matricule,))
                contract_data = cursor.fetchone()

                if not employee_data:
                    Messagebox.show_error("Employé non trouvé", "Erreur")
                    return

                cursor.execute("DELETE FROM contrats WHERE matricule = %s", (matricule,))
                cursor.execute("DELETE FROM employes WHERE matricule = %s", (matricule,))
                self.conn.commit()

                self.undo_stack.append((
                    "employee_delete", matricule,
                    {
                        'matricule': employee_data[0], 'nom': employee_data[1], 'prenom': employee_data[2],
                        'genre': employee_data[3], 'date_naissance': employee_data[4],
                        'lieu_naissance': employee_data[5],
                        'adresse': employee_data[6], 'ville': employee_data[7], 'cin': employee_data[8],
                        'date_cin': employee_data[9], 'lieu_cin': employee_data[10], 'poste': employee_data[11],
                        'email': employee_data[12], 'telephone': employee_data[13]
                    },
                    {
                        'id': contract_data[0], 'type_contrat': contract_data[2], 'date_creation': contract_data[3],
                        'date_debut': contract_data[4], 'date_fin': contract_data[5], 'salaire_base': contract_data[6],
                        'prime': contract_data[7], 'salary_type': contract_data[8], 'texte_contrat': contract_data[9]
                    } if contract_data else None
                ))

                self.load_data()
                self.clear_search()
                self.status_var.set(f"Employé {matricule} supprimé")
        except mysql.connector.Error as e:
            Messagebox.show_error(f"Erreur de suppression: {str(e)}", "Erreur")

    def undo_action(self):
        if not self.undo_stack:
            Messagebox.show_info("Aucune action à annuler", "Information")
            return

        action_type, matricule, employee_data, contract_data = self.undo_stack.pop()
        try:
            with self.conn.cursor() as cursor:
                if action_type == "employee_delete":
                    cursor.execute('''
                                   INSERT INTO employes (matricule, nom, prenom, genre, date_naissance, lieu_naissance,
                                                         adresse, ville, cin, date_cin, lieu_cin, poste, email,
                                                         telephone)
                                   VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                   ''', (
                                       employee_data['matricule'], employee_data['nom'], employee_data['prenom'],
                                       employee_data['genre'],
                                       employee_data['date_naissance'], employee_data['lieu_naissance'],
                                       employee_data['adresse'],
                                       employee_data['ville'], employee_data['cin'], employee_data['date_cin'],
                                       employee_data['lieu_cin'],
                                       employee_data['poste'], employee_data['email'], employee_data['telephone']
                                   ))
                    if contract_data:
                        cursor.execute('''
                                       INSERT INTO contrats (id, matricule, type_contrat, date_creation, date_debut,
                                                             date_fin, salaire_base, prime, salary_type, texte_contrat)
                                       VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                                       ''', (
                                           contract_data['id'], matricule, contract_data['type_contrat'],
                                           contract_data['date_creation'],
                                           contract_data['date_debut'], contract_data['date_fin'],
                                           contract_data['salaire_base'],
                                           contract_data['prime'], contract_data['salary_type'],
                                           contract_data['texte_contrat']
                                       ))
                    self.status_var.set(f"Suppression de {matricule} annulée")
                elif action_type == "employee_update":
                    cursor.execute(f"UPDATE employes SET {employee_data} = %s WHERE matricule = %s",
                                   (contract_data, matricule))
                    self.status_var.set(f"Mise à jour de {employee_data} pour {matricule} annulée")
                elif action_type == "contract_update":
                    cursor.execute(f"UPDATE contrats SET {employee_data} = %s WHERE id = %s",
                                   (contract_data, matricule))
                    self.status_var.set(f"Mise à jour du contrat pour {matricule} annulée")

                self.conn.commit()
                self.load_data()
                if self.current_employee and self.current_employee['matricule'] == matricule:
                    self.search_employee()
        except mysql.connector.Error as e:
            Messagebox.show_error(f"Erreur lors de l'annulation: {str(e)}", "Erreur")

    def show_help(self):
        Messagebox.show_info(
            """Application de Gestion des Contrats
    Version: 1.0
    Fonctionnalités:
    - Ajouter, modifier, supprimer des employés
    - Générer des contrats CDD/CDI en arabe
    - Exporter les contrats en Word ou ZIP
    - Rechercher et filtrer les employés
    - Modifier les données directement dans le tableau
    - Annuler la dernière action (suppression ou modification)

    Pour plus d'aide, contactez le support technique.""",
            "Aide"
        )

    def show_about(self):
        Messagebox.show_info(
            """Application de Gestion des Contrats
    Développée par: Imbert Mnif
    Version: 1.0
    © Imbert Mnif. Tous droits réservés.""",
            "À propos"
        )

    def __del__(self):
        self.stop_alert_timer()
        self.stop_alert_sound()
        if hasattr(self, 'conn') and self.conn.is_connected():
            self.conn.close()


    # Puis ta fonction :
    def check_expiring_contracts(self, force_notification=False):
        try:
            today = datetime.datetime.now().date()

            cursor = self.conn.cursor()
            query = '''
                    SELECT e.matricule,
                           e.nom,
                           e.prenom,
                           c.date_fin,
                           DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) AS jours_restants
                    FROM employes e
                             JOIN contrats c ON e.matricule = c.matricule
                    WHERE c.type_contrat = 'CDD'
                      AND c.date_fin IS NOT NULL
                      AND DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) BETWEEN 0 AND 30
                    ORDER BY jours_restants \
                    '''
            cursor.execute(query, (today, today))
            expiring_contracts = cursor.fetchall()
            cursor.close()

            new_alerts = []

            if expiring_contracts:
                for contract in expiring_contracts:
                    matricule, nom, prenom, date_fin, jours_restants = contract
                    contract_key = f"{matricule}_{date_fin.strip().lower()}"

                    if contract_key not in self.notified_contracts:
                        new_alerts.append(contract)
                        self.notified_contracts.add(contract_key)

            if new_alerts and force_notification:
                message = "⚠️ ALERTE : Contrats CDD expirant bientôt ⚠️\n\n"
                for contract in new_alerts:
                    matricule, nom, prenom, date_fin, jours_restants = contract
                    message += f"• {nom} {prenom} (Matricule: {matricule}) - "
                    message += f"Expire le {date_fin} (dans {jours_restants} jours)\n"
                print(f"Notification envoyée pour {len(new_alerts)} contrats")  # DEBUG
                self.play_alert_sound()
                messagebox.showwarning("Alerte Contrats", message, parent=self.root)
                self.stop_alert_sound()
                self.status_var.set(f"⚠ {len(new_alerts)} nouveaux contrats expirent dans ≤30 jours")

        except Exception as e:
            print(f"Erreur lors de la vérification des contrats: {e}")

    def play_alert_sound(self):
        """Joue le son d'alerte toutes les 2 secondes jusqu'à ce que l'utilisateur clique sur OK"""
        if self.sound_enabled and hasattr(self, 'sound_file') and self.sound_file and os.path.exists(self.sound_file):
            try:
                def play_loop():
                    while not hasattr(self, 'alert_stopped') or not self.alert_stopped:
                        winsound.PlaySound(self.sound_file, winsound.SND_FILENAME | winsound.SND_ASYNC)
                        time.sleep(2)  # Répéter toutes les 2 secondes

                self.alert_thread = threading.Thread(target=play_loop)
                self.alert_thread.daemon = True
                self.alert_thread.start()
            except Exception as e:
                print(f"Erreur de lecture du son d'alerte: {e}")
                try:
                    def play_system_loop():
                        while not hasattr(self, 'alert_stopped') or not self.alert_stopped:
                            winsound.PlaySound("SystemExclamation", winsound.SND_ALIAS | winsound.SND_ASYNC)
                            time.sleep(2)

                    self.alert_thread = threading.Thread(target=play_system_loop)
                    self.alert_thread.daemon = True
                    self.alert_thread.start()
                except:
                    pass

    def stop_alert_sound(self):
        """Arrête le son d'alerte"""
        try:
            self.alert_stopped = True
            winsound.PlaySound(None, 0)
            if hasattr(self, 'alert_thread'):
                self.alert_thread.join(timeout=0.1)
        except Exception as e:
            print(f"Erreur lors de l'arrêt du son: {e}")

    def stop_alert_timer(self):
        if self.alert_timer:
            self.root.after_cancel(self.alert_timer)

    def update_summary(self):
        self.check_expiring_contracts(force_notification=True)  # Première vérification

        try:
            with self.conn.cursor() as cursor:
                # Compter le nombre total d'employés uniques
                cursor.execute("SELECT COUNT(DISTINCT matricule) FROM employes")
                total_employees = cursor.fetchone()[0]

                # Compter les contrats actifs (dernier contrat de chaque employé)
                cursor.execute('''
                               SELECT COUNT(CASE WHEN c.type_contrat = 'CDD' THEN 1 END) AS cdd_count,
                                      COUNT(CASE WHEN c.type_contrat = 'CDI' THEN 1 END) AS cdi_count,
                                      AVG(c.salaire_base)                                AS avg_salary
                               FROM (SELECT matricule, MAX(id) AS max_id
                                     FROM contrats
                                     GROUP BY matricule) AS latest
                                        JOIN contrats c ON c.id = latest.max_id
                               ''')
                stats = cursor.fetchone()
                cdd_count = stats[0] or 0
                cdi_count = stats[1] or 0
                avg_salary = stats[2] or 0

                # Compter les CDD expirant bientôt (dans les 30 jours)
                today = datetime.datetime.now().date()
                cursor.execute('''
                               SELECT COUNT(*)
                               FROM (SELECT c.matricule, MAX(c.id) AS max_id
                                     FROM contrats c
                                     WHERE c.type_contrat = 'CDD'
                                     GROUP BY c.matricule) AS latest
                                        JOIN contrats c ON c.id = latest.max_id
                               WHERE c.date_fin IS NOT NULL
                                 AND DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) <= 30
                                 AND DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) >= 0
                               ''', (today, today))
                expiring_soon = cursor.fetchone()[0] or 0

                # Obtenir la liste détaillée
                cursor.execute('''
                               SELECT e.matricule,
                                      e.nom,
                                      e.prenom,
                                      c.date_fin,
                                      DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) AS jours_restants
                               FROM employes e
                                        JOIN contrats c ON e.matricule = c.matricule
                               WHERE c.type_contrat = 'CDD'
                                 AND c.date_fin IS NOT NULL
                                 AND DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) BETWEEN 0 AND 30
                               ORDER BY jours_restants
                               ''', (today, today))
                expiring_contracts = cursor.fetchall()

                # Affichage du message d'alerte
                if expiring_soon > 0:
                    self.play_alert_sound()  # Lancer le son avant l'affichage du message
                    message = "⚠️ ALERTE : Contrats CDD expirant bientôt ⚠️\n\n"
                    for contract in expiring_contracts:
                        matricule, nom, prenom, date_fin, jours_restants = contract
                        message += f"• {nom} {prenom} (Matricule: {matricule}) - "
                        message += f"Expire le {date_fin} (dans {jours_restants} jours)\n"
                    Messagebox.show_warning(message, "Alerte Contrats", parent=self.root)
                    self.stop_alert_sound()  # Arrêter le son après la fermeture du message
                    self.status_var.set(f"⚠ {expiring_soon} contrats expirent dans ≤30 jours")
                else:
                    self.status_var.set("Aucun contrat expirant bientôt")

            # Mettre à jour l'interface
            self.total_label.config(text=f"Employés: {total_employees}")
            self.cdd_label.config(text=f"CDD Actifs: {cdd_count}")
            self.cdi_label.config(text=f"CDI Actifs: {cdi_count}")
            self.salary_label.config(text=f"Salaire Moyen: {avg_salary:.2f} TND")

        except mysql.connector.Error as e:
            print(f"Erreur lors de la mise à jour du résumé: {e}")
            self.status_var.set("Erreur lors de la mise à jour des statistiques")

    def background_alert_service(self):
        """Service en arrière-plan pour vérifier les contrats proches d'expiration et envoyer des notifications système"""
        try:
            if not hasattr(self, 'conn') or not self.conn.is_connected():
                self.conn = mysql.connector.connect(**self.db_config)
            today = datetime.datetime.now().date()
            with self.conn.cursor() as cursor:
                cursor.execute('''
                               SELECT e.matricule,
                                      e.nom,
                                      e.prenom,
                                      c.date_fin,
                                      DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) AS jours_restants
                               FROM employes e
                                        JOIN contrats c ON e.matricule = c.matricule
                               WHERE c.type_contrat = 'CDD'
                                 AND c.date_fin IS NOT NULL
                                 AND DATEDIFF(STR_TO_DATE(c.date_fin, '%d/%m/%Y'), %s) BETWEEN 0 AND 30
                               ORDER BY jours_restants
                               ''', (today, today))
                expiring_contracts = cursor.fetchall()

                if expiring_contracts:
                    current_check = []
                    for contract in expiring_contracts:
                        matricule, nom, prenom, date_fin, jours_restants = contract
                        contract_data = f"{matricule}{date_fin}"
                        contract_hash = hashlib.md5(contract_data.encode()).hexdigest()

                        if matricule not in self.contract_hash or self.contract_hash[matricule] != contract_hash:
                            self.contract_hash[matricule] = contract_hash
                            title = "Alerte Contrat"
                            message = f"Contrat de {nom} {prenom} (Matricule: {matricule}) expire le {date_fin} (dans {jours_restants} jours)"
                            try:
                                notification.notify(
                                    title=title,
                                    message=message,
                                    app_name="Gestion des Contrats",
                                    timeout=10
                                )
                            except Exception as e:
                                print(f"Erreur lors de l'envoi de la notification: {e}")

                        current_check.append((matricule, date_fin))

                    self.last_contract_check = tuple(current_check)
                    self.status_var.set(f"⚠ {len(expiring_contracts)} contrats expirent dans ≤30 jours")

                    # Afficher l'alerte initiale
                    self.check_expiring_contracts(True)

                else:
                    self.status_var.set("Aucun contrat expirant bientôt")

                # Planifier la prochaine vérification
                self.alert_timer = self.root.after(self.check_interval, self.background_alert_service)

        except Exception as e:
            print(f"Erreur dans le service d'alerte en arrière-plan: {e}")
            self.alert_timer = self.root.after(self.check_interval, self.background_alert_service)

    def get_widget_value(self, widget):
        """Obtient la valeur de n'importe quel widget"""
        if isinstance(widget, (ttk.Entry, ttk.Combobox)):
            return widget.get().strip()
        elif isinstance(widget, DateEntry):
            return widget.entry.get().strip()
        elif hasattr(widget, 'get'):
            return widget.get().strip()
        return ""


    def create_contract_doc(self, matricule, texte):
        # Récupérer les données de l'employé et du contrat depuis la base
        with self.conn.cursor() as cursor:
            cursor.execute('''
                SELECT e.*, c.type_contrat, c.date_debut, c.date_fin, 
                       c.salaire_base, c.prime, c.salary_type
                FROM employes e
                JOIN contrats c ON e.matricule = c.matricule
                WHERE e.matricule = %s
                ORDER BY c.id DESC LIMIT 1
            ''', (matricule,))
            result = cursor.fetchone()

            if not result:
                return None

            # Données employé
            employee_data = {
                'matricule': result[0],
                'nom': result[1],
                'prenom': result[2],
                'genre': result[3],
                'date_naissance': result[4] or "غير محدد",
                'lieu_naissance': result[5] or "غير محدد",
                'ville': result[7] or "غير محدد",
                'cin': result[8] or "غير محدد",
                'date_cin': result[9] or "غير محدد",
                'lieu_cin': result[10] or "غير محدد",
                'poste': result[11] or "غير محدد"
            }

            # Données contrat
            contract_data = {
                'type_contrat': result[14],
                'date_debut': result[15],
                'date_fin': result[16] or "غير محدد",
                'salaire_base': result[17],
                'prime': result[18],
                'salary_type': "الساعة" if result[19] == "hourly" else "الشهر"
            }

        # Calcul de la durée pour CDD
        duree_txt = "غير محددة"
        if contract_data['type_contrat'] == "CDD" and result[16]:
            try:
                debut = datetime.datetime.strptime(contract_data['date_debut'], "%d/%m/%Y")
                fin = datetime.datetime.strptime(contract_data['date_fin'], "%d/%m/%Y")
                duree_jours = (fin - debut).days
                mois = duree_jours // 30
                jours = duree_jours % 30
                duree_txt = f"{mois} شهرا" + (f" و {jours} يوما" if jours else "")
            except ValueError:
                pass

        # Sélection du modèle approprié
        if contract_data['type_contrat'] == "CDD":
            if employee_data['genre'] == "féminin":
                template = self.CDD_FEMININ
            else:
                template = self.CDD_MASCULIN
        else:  # CDI
            if employee_data['genre'] == "féminin":
                template = self.CDI_FEMININ
            else:
                template = self.CDI_MASCULIN

        # Remplissage du modèle
        contrat_text = template
        contrat_text = contrat_text.replace("{{Mat}}", employee_data['matricule'])
        contrat_text = contrat_text.replace("{{Titre}}", "الآنسة" if employee_data['genre'] == "féminin" else "السيد")
        contrat_text = contrat_text.replace("{{Prénom}}", employee_data['prenom'])
        contrat_text = contrat_text.replace("{{Nom}}", employee_data['nom'])
        contrat_text = contrat_text.replace("{{DNAIS}}", employee_data['date_naissance'])
        contrat_text = contrat_text.replace("{{LNAIS}}", employee_data['lieu_naissance'])
        contrat_text = contrat_text.replace("{{Ville}}", employee_data['ville'])
        contrat_text = contrat_text.replace("{{NCIN}}", employee_data['cin'])
        contrat_text = contrat_text.replace("{{DCIN}}", employee_data['date_cin'])
        contrat_text = contrat_text.replace("{{LCIN}}", employee_data['lieu_cin'])
        contrat_text = contrat_text.replace("{{Poste}}", employee_data['poste'])
        contrat_text = contrat_text.replace("{{DPERIODE}}", contract_data['date_debut'])
        contrat_text = contrat_text.replace("{{FPERIODE}}", contract_data['date_fin'])
        contrat_text = contrat_text.replace("{{DUREE}}", duree_txt)
        contrat_text = contrat_text.replace("{{SBASE}}", str(contract_data['salaire_base']))
        contrat_text = contrat_text.replace("{{PRIME}}", str(contract_data['prime']))
        contrat_text = contrat_text.replace("{{MPAIE}}", contract_data['salary_type'])
        contrat_text = contrat_text.replace("{{DATE_CONTRAT}}", datetime.datetime.now().strftime('%d/%m/%Y'))

        # Création du document Word
        doc = Document()
        section = doc.sections[0]
        section.left_margin = section.right_margin = Inches(8 / 25.4)  # 8mm
        section.top_margin = section.bottom_margin = Inches(5 / 25.4)  # 5mm
        section.is_right_to_left = True

        # Tableau principal
        main_table = doc.add_table(rows=1, cols=3)
        main_table.columns[0].width = Inches(15 / 25.4)
        main_table.columns[1].width = Inches(80 / 25.4)
        main_table.columns[2].width = Inches(40 / 25.4)
        main_table.rows[0].height = Inches(15 / 25.4)
        main_table.style = 'Table Grid'

        # Cellule Logo
        logo_cell = main_table.cell(0, 0)
        try:
            logo_para = logo_cell.add_paragraph()
            logo_run = logo_para.add_run()
            logo_run.add_picture(self.logo_path, width=Inches(12 / 25.4), height=Inches(12 / 25.4))
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Erreur lors du chargement du logo : {e}")
            logo_cell.text = ""

        # Cellule Titre
        title_cell = main_table.cell(0, 1)
        title_para = title_cell.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        french_title = title_para.add_run("FORMULAIRE\n")
        french_title.bold = True
        french_title.font.name = "Helvetica"
        french_title.font.rtl = True

        arabic_title = title_para.add_run(
            f"عقد شغل لمدة {'غير محددة' if contract_data['type_contrat'] == 'CDI' else 'محدودة'}\n")
        arabic_title.bold = True
        arabic_title.font.name = "Arial"

        arabic_title.font.rtl = True

        # Tableau d'informations
        info_cell = main_table.cell(0, 2)
        info_table = info_cell.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.columns[0].width = info_table.columns[1].width = Inches(15 / 25.4)

        info_data = [
            ("Réf.", f"FO-RH-0{4 if contract_data['type_contrat'] == 'CDI' else 3}"),
            ("Date", datetime.datetime.now().strftime('%d/%m/%Y')),
            ("Version", "01"),
            ("Page", "1/1")
        ]

        for row_idx, (label, value) in enumerate(info_data):
            label_cell = info_table.cell(row_idx, 0)
            label_para = label_cell.add_paragraph(label)
            label_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            label_run = label_para.runs[0]
            label_run.font.name = "Helvetica"
            label_run.font.size = Pt(7)
            label_run.font.color.rgb = RGBColor(100, 100, 100)
            label_run.font.rtl = True

            value_cell = info_table.cell(row_idx, 1)
            value_para = value_cell.add_paragraph(value)
            value_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            value_run = value_para.runs[0]
            value_run.font.name = "Helvetica"
            value_run.font.size = Pt(7)
            value_run.font.rtl = True

        # Matricule
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        matricule_para = doc.add_paragraph()
        matricule_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        matricule_run = matricule_para.add_run(f"Matricule: {employee_data['matricule']}")
        matricule_run.bold = True
        matricule_run.font.name = "Arial"
        matricule_run.font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Ajout du texte du contrat
        for paragraph in contrat_text.strip().split('\n'):
            if paragraph.strip():
                body_para = doc.add_paragraph()
                body_para.paragraph_format.space_after = Pt(2)
                body_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                body_run = body_para.add_run(paragraph.strip())
                body_run.font.name = "Arial"
                body_run.font.size = Pt(9)
                body_run.font.rtl = True

        return doc


    def save_and_generate(self):
        """Enregistre un nouvel employé et son contrat, puis génère et affiche le texte du contrat."""
        try:
            # 1. Récupération des données des widgets
            employee_data = {
                'matricule': self.get_widget_value(self.entries['matricule']),
                'nom': self.get_widget_value(self.entries['nom']),
                'prenom': self.get_widget_value(self.entries['prenom']),
                'genre': self.variables["genre"].get(),
                'date_naissance': self.get_widget_value(self.entries['date_naissance']) or None,
                'lieu_naissance': self.get_widget_value(self.entries['lieu_naissance']) or None,
                'adresse': self.get_widget_value(self.entries['adresse']) or None,
                'ville': self.get_widget_value(self.entries['ville']) or None,
                'cin': self.get_widget_value(self.entries['cin']) or None,
                'date_cin': self.get_widget_value(self.entries['date_cin']) or None,
                'lieu_cin': self.get_widget_value(self.entries['lieu_cin']) or None,
                'poste': self.get_widget_value(self.entries['poste']) or None,
                'email': self.get_widget_value(self.entries['email']) or None,
                'telephone': self.get_widget_value(self.entries['telephone']) or None
            }

            contract_data = {
                'matricule': employee_data['matricule'],
                'type_contrat': self.variables["contract_type"].get(),
                'date_debut': self.get_widget_value(self.contract_entries['date_debut']),
                'date_fin': self.get_widget_value(self.contract_entries['date_fin']) if self.variables[
                                                                                            "contract_type"].get() == "CDD" else None,
                'salaire_base': float(self.get_widget_value(self.contract_entries['salaire'])),
                'prime': float(self.get_widget_value(self.contract_entries['prime'])),
                'salary_type': self.variables["salary_type"].get()
            }

            # 2. Validation des champs obligatoires
            required_fields = {
                'matricule': employee_data['matricule'],
                'nom': employee_data['nom'],
                'prenom': employee_data['prenom'],
                'date_debut': contract_data['date_debut'],
                'salaire_base': contract_data['salaire_base'],
                'prime': contract_data['prime']
            }

            for field, value in required_fields.items():
                if not value or (isinstance(value, str) and value.isspace()):
                    Messagebox.show_error(f"Le champ {field} est obligatoire.", "Erreur de validation")
                    return

            # 3. Validation des formats
            try:
                datetime.datetime.strptime(contract_data['date_debut'], "%d/%m/%Y")
                if contract_data['date_fin']:
                    datetime.datetime.strptime(contract_data['date_fin'], "%d/%m/%Y")
            except ValueError:
                Messagebox.show_error("Format de date invalide (JJ/MM/AAAA attendu)", "Erreur de validation")
                return

            # 4. Vérification de l'unicité du matricule
            with self.conn.cursor() as cursor:
                cursor.execute("SELECT matricule FROM employes WHERE matricule = %s", (employee_data['matricule'],))
                if cursor.fetchone():
                    Messagebox.show_error("Ce matricule existe déjà.", "Erreur de validation")
                    return

            # 5. Calcul de la durée pour CDD
            duree_txt = "غير محددة"
            if contract_data['type_contrat'] == "CDD" and contract_data['date_fin']:
                try:
                    debut = datetime.datetime.strptime(contract_data['date_debut'], "%d/%m/%Y")
                    fin = datetime.datetime.strptime(contract_data['date_fin'], "%d/%m/%Y")
                    duree_jours = (fin - debut).days
                    mois = duree_jours // 30
                    jours = duree_jours % 30
                    duree_txt = f"{mois} شهرا" + (f" و {jours} يوما" if jours else "")

                    # Vérification durée minimale (30 jours)
                    if duree_jours < 30:
                        Messagebox.show_warning(
                            "Attention: La durée du CDD est inférieure à 30 jours. "
                            "Veuillez vérifier les dates de début et fin.",
                            "Durée minimale non respectée"
                        )
                except ValueError:
                    pass

            # 6. Sélection du modèle de contrat approprié
            if contract_data['type_contrat'] == "CDD":
                if employee_data['genre'] == "féminin":
                    template = self.CDD_FEMININ
                else:
                    template = self.CDD_MASCULIN
            else:  # CDI
                if employee_data['genre'] == "féminin":
                    template = self.CDI_FEMININ
                else:
                    template = self.CDI_MASCULIN

            # 7. Remplacement des variables dans le modèle
            contrat_text = template
            replacements = {
                '{{Mat}}': employee_data['matricule'],
                '{{Titre}}': "الآنسة" if employee_data['genre'] == "féminin" else "السيد",
                '{{Prénom}}': employee_data['prenom'],
                '{{Nom}}': employee_data['nom'],
                '{{DNAIS}}': employee_data['date_naissance'] or "غير محدد",
                '{{LNAIS}}': employee_data['lieu_naissance'] or "غير محدد",
                '{{Ville}}': employee_data['ville'] or "غير محدد",
                '{{NCIN}}': employee_data['cin'] or "غير محدد",
                '{{DCIN}}': employee_data['date_cin'] or "غير محدد",
                '{{LCIN}}': employee_data['lieu_cin'] or "غير محدد",
                '{{Poste}}': employee_data['poste'] or "غير محدد",
                '{{DPERIODE}}': contract_data['date_debut'],
                '{{FPERIODE}}': contract_data['date_fin'] or "غير محدد",
                '{{DUREE}}': duree_txt,
                '{{SBASE}}': str(contract_data['salaire_base']),
                '{{PRIME}}': str(contract_data['prime']),
                '{{MPAIE}}': "الساعة" if contract_data['salary_type'] == "hourly" else "الشهر",
                '{{DATE_CONTRAT}}': datetime.datetime.now().strftime('%d/%m/%Y')
            }

            for placeholder, value in replacements.items():
                contrat_text = contrat_text.replace(placeholder, value)

            # 8. Enregistrement dans la base de données
            with self.conn.cursor() as cursor:
                # Insertion de l'employé
                cursor.execute('''
                               INSERT INTO employes (matricule, nom, prenom, genre, date_naissance, lieu_naissance,
                                                     adresse, ville, cin, date_cin, lieu_cin, poste, email, telephone)
                               VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                               ''', (
                                   employee_data['matricule'], employee_data['nom'], employee_data['prenom'],
                                   employee_data['genre'], employee_data['date_naissance'],
                                   employee_data['lieu_naissance'],
                                   employee_data['adresse'], employee_data['ville'], employee_data['cin'],
                                   employee_data['date_cin'], employee_data['lieu_cin'], employee_data['poste'],
                                   employee_data['email'], employee_data['telephone']
                               ))

                # Insertion du contrat
                cursor.execute('''
                               INSERT INTO contrats (matricule, type_contrat, date_creation, date_debut, date_fin,
                                                     salaire_base, prime, salary_type, texte_contrat)
                               VALUES (%s, %s, NOW(), %s, %s, %s, %s, %s, %s)
                               ''', (
                                   contract_data['matricule'], contract_data['type_contrat'],
                                   contract_data['date_debut'], contract_data['date_fin'],
                                   contract_data['salaire_base'], contract_data['prime'],
                                   contract_data['salary_type'], contrat_text
                               ))

                self.conn.commit()

            # 9. Affichage du contrat généré
            self.current_employee = employee_data
            self.current_employee['texte_contrat'] = contrat_text
            self.contract_text.delete(1.0, tk.END)
            self.contract_text.insert(tk.END, contrat_text)
            self.notebook.select(2)  # Sélectionne l'onglet Contrat

            # 10. Mise à jour de l'interface
            self.load_data()
            self.clear_form()
            self.status_var.set(
                f"Contrat {contract_data['type_contrat']} enregistré pour {employee_data['prenom']} {employee_data['nom']}")
            Messagebox.show_info("Enregistrement et génération réussis.", "Succès")

        except mysql.connector.Error as e:
            self.conn.rollback()
            Messagebox.show_error(f"Erreur de base de données : {str(e)}", "Erreur")
        except ValueError as e:
            Messagebox.show_error(f"Erreur de valeur : {str(e)}", "Erreur")
        except Exception as e:
            Messagebox.show_error(f"Erreur inattendue : {str(e)}", "Erreur")

if __name__ == "__main__":
    app = ContratApplication(ttk.Window())
    app.root.mainloop()